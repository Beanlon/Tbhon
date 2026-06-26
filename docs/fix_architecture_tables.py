"""Fix table placement and captions in JAM_Architecture Paper.docx Section IV."""
from __future__ import annotations

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC = r"C:\Users\Mika\Downloads\JAM_Architecture Paper.docx"


def insert_table_after(paragraph: Paragraph, rows: int, cols: int) -> Table:
    doc = paragraph.part.document
    table = doc.add_table(rows=rows, cols=cols)
    paragraph._p.addnext(table._tbl)
    return table


def set_cell(table: Table, row: int, col: int, value: str) -> None:
    table.rows[row].cells[col].text = value


def fill_table(table: Table, headers: list[str], data: list[list[str]]) -> None:
    for col, header in enumerate(headers):
        set_cell(table, 0, col, header)
    for row_idx, row in enumerate(data, start=1):
        for col_idx, value in enumerate(row):
            set_cell(table, row_idx, col_idx, value)


def remove_table(table: Table) -> None:
    table._element.getparent().remove(table._element)


def caption_after(doc: Document, prefix: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise RuntimeError(f"Caption not found: {prefix}")


def main() -> None:
    doc = Document(SRC)

    # Remove orphaned tables at end (keep first table = cross-fold at TABLE II).
    while len(doc.tables) > 1:
        remove_table(doc.tables[-1])

    # TABLE I ablation
    p1 = caption_after(doc, "TABLE I")
    p1.text = "TABLE I\nCOUGH CLASSIFIER ABLATION ON FOLD 0 (HELD-OUT TEST, n = 2,606)"
    t1 = insert_table_after(p1, 3, 7)
    fill_table(
        t1,
        ["Model", "Fold", "Test n", "Accuracy", "Precision", "Recall", "Macro-F1"],
        [
            ["Convolutional network (baseline)", "0", "2,606", "50.6%", "60.4%", "60.7%", "50.6%"],
            ["Hybrid CNN+GBM (ablation)", "0", "2,606", "73.9%", "68.3%", "68.4%", "68.3%"],
        ],
    )

    p2 = caption_after(doc, "TABLE II")
    p2.text = "TABLE II\nHYBRID COUGH CLASSIFIER CROSS-FOLD TEST PERFORMANCE"

    p3 = caption_after(doc, "TABLE III")
    p3.text = "TABLE III\nPRODUCTION UNIMODAL CLASSIFIER PERFORMANCE ON HELD-OUT TEST DATA"
    t3 = insert_table_after(p3, 4, 6)
    fill_table(
        t3,
        ["Production Model", "Test n", "Accuracy", "Macro Precision", "Macro Recall", "Macro-F1"],
        [
            ["Hybrid cough CNN+GBM (fold 1)", "2,606", "75.8%", "71.7%", "67.9%", "69.1%"],
            ["ResNet-18 AFB binary (sputum)", "216", "94.9%", "62.6%", "73.0%", "66.3%"],
            ["AFB+ sensitivity (sputum)", "210 positives", "N/A", "98.5%", "96.2%", "97.1%"],
        ],
    )

    p4 = caption_after(doc, "TABLE IV")
    p4.text = "TABLE IV\nMULTIMODAL FUSION MODALITY RELIABILITY WEIGHTS"
    t4 = insert_table_after(p4, 4, 2)
    fill_table(
        t4,
        ["Modality", "Reliability Weight"],
        [
            ["Symptom and exposure checklist", "0.85"],
            ["Cough machine learning probability", "1.00"],
            ["Sputum machine learning probability", "0.70"],
        ],
    )

    # Normalize cross-fold table (doc.tables[0]) mean row dash
    t0 = doc.tables[0]
    set_cell(t0, 4, 1, "N/A")

    doc.save(SRC)
    print("Fixed table placement in", SRC)


if __name__ == "__main__":
    main()
