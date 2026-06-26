"""Generate TBhon pitch deck speaker script — full read-aloud prose for 3 members."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "TBhon_Pitch_Deck_Script.docx"
OUT_ALT = ROOT / "docs" / "TBhon_Pitch_Deck_Script_generated.docx"


MEMBER_1_SCRIPT = """\
[Slide 1 — Title]

What if finding TB didn't need a lab — just a cough, a phone, and 5 minutes?

We're TBhon — reinventing TB risk detection at the edge. We're building the first line of defense before the microscope, before GeneXpert, before it's too late.


[Slide 2 — The Problem, three photos]

TB is a silent epidemic — and our screening gap is killing people.

Picture a farmer in a rural barangay. He's been coughing for 3 weeks. Night sweats. Weight loss. The nearest RHU is hours away. The GeneXpert machine? Maybe one per city — and the sign on the wall says 2 days' wait.

So what happens? They wait. They guess. And they spread it to their family — the people sitting right beside them at home.


[Slide 3 — Triage flowchart]

And here's the part nobody talks about enough: the real problem isn't treatment. It's triage.

Who needs urgent referral today — and who can go home? Right now that decision is made with paper forms, gut feel, and overloaded health workers stuck in the bottleneck between the waiting room and the lab.


[Slide 4 — Title recap]

That's the gap TBhon was built to close.


[Slide 5 — Stats: 1.2M / 60% / 18%]

The numbers back this up. Tuberculosis remains one of the world's deadliest infectious killers — 1.2M deaths every year. 60% of cases can present with zero classic symptoms, driving silent transmission. And frontline health workers face an 18% annual infection rate when triage happens too late in crowded waiting rooms — like the one in this photo.


[Slide 6 — Three signals]

So what does TBhon do?

TBhon answers one question in minutes: Should this person be referred for confirmatory TB testing — right now?

This is not a diagnosis. It is not a replacement for GeneXpert. It is a smart triage decision backed by three signals: how they sound — cough audio plus AI; what their sputum shows — microscopy image plus AI; and what symptoms they report — our 11-question clinical checklist.

One app. One result: Low, Moderate, or High risk.

I'll hand over to [Member 2 name] to walk you through how we built it.\
"""

MEMBER_2_SCRIPT = """\
[Slide 7 — How our features bridge the gap]

Thank you. Here's what that looks like in the actual app.

On the left you see a fused Low Risk result — sputum sample from the server and a fusion model breakdown showing 30.8% combined probability. In the center, IoT sputum capture from our booth device. On the right, a cough recording with quality feedback before upload.

Under the hood we use Mel-spectrogram CNN analysis for cough audio, ResNet18 for Ziehl–Neelsen smear images, and PyTorch multimodal scoring fused into one risk band. Capture at the booth. Inference in the cloud. Decision on the phone.


[Slide 8 — Pre-screening comparison table]

Why does that middle layer matter? Symptom self-checks are accessible — but they miss more than half of active cases. CAD X-ray and GeneXpert-class tools are accurate — but they're locked behind clinic walls and expensive infrastructure.

TBhon sits in the middle: high accessibility on smartphones and IoT, a reliable pre-screening target, and extremely low infrastructure requirement. We support diagnosis — we do not substitute it.


[Slide 9 — Key Features]

Three pillars in one workflow. Multimodal TB screening — checklist, cough audio, and sputum image in one mobile session. AI-assisted analysis — real-time cough quality checks, ML classification, and unified risk scoring. And IoT bench support — our ESP32 device with BLE and Wi-Fi provisioning for assisted cough and sputum collection.


[Slide 10 — Technology Stack]

Our stack is production-grade. Frontend: Expo Router, React Native, TypeScript, NativeWind. Backend: Node.js, Express, Prisma ORM, JWT auth, deployed on DigitalOcean. Machine learning: Python, FastAPI, Uvicorn, PyTorch, and scikit-learn for the hybrid cough stack. This is not a localhost demo — it runs on real cloud infrastructure.


[Slide 11 — Methods: Three Models, One Decision]

Three models. One decision.

Row one: the symptom checklist — 11 clinical questions scored with a hand-tuned logistic model. It catches red flags — cough of 2+ weeks, hemoptysis, night sweats, weight loss, TB contact.

Row two: Cough AI — Kaggle TB audio dataset, 2,606 held-out test samples, hybrid CNN plus GBM on Mel-spectrograms.

Row three: Sputum AI — microscopy smears, YOLO-annotated dataset, 216 test images, ResNet18 AFB classifier.


[Slide 12 — Cough pipeline flowchart]

The cough pipeline in detail: raw wav audio, resampled to 16 kHz, cropped or padded to 4 seconds. Converted to a Mel spectrogram and fed through LegacySmallAudioCNN for feature extraction. Those features go into gradient boosting — XGBoost or LightGBM. CNN and GBM probabilities are blended with a tuned threshold for the final TB versus No-TB output.


[Slide 13 — Sputum dataset workflow]

On the sputum side: collect de-identified microscopy images, annotate AFB with YOLO bounding boxes, count bacilli per field, grade by the WHO scale, then split at the patient level — 70% train, 15% validation, 15% test — so there's no data leakage. Quality control removes blur, empty fields, and bad annotations.


[Slide 14 — Weighted log-odds fusion]

We don't trust any single signal alone. Checklist, cough, and sputum probabilities combine through weighted log-odds fusion — weights 0.85 for checklist, 1.00 for cough, 0.70 for sputum.

The app outputs three risk bands: Low below 38%, Moderate from 38% to 62%, High at 62% and above. Note: the graphic on this slide shows older four-band labels — follow what I'm saying, not the outdated numbers on the image.

Clinical safety floors raise the fused probability for hemoptysis, high checklist concern, and confident AFB-positive smear. Screening triage only — not a diagnosis.


[Slide 15 — ESP32 hardware booth]

We didn't stop at software. We built the booth. INMP441 microphone for clean cough capture in the field. OV5640 camera for microscopic sputum imaging. ESP32-S3 with BLE and Wi-Fi provisioning — staff sets it up from the app in seconds.

Device captures. Cloud thinks. App decides.

I'll pass to [Member 3 name] for validation and where we're headed next.\
"""

MEMBER_3_SCRIPT = """\
[Slide 16 — We didn't just build it. We tested it.]

Thank you. We didn't just build it. We tested it.

On held-out test sets: our cough model — hybrid CNN plus GBM — reached 75.75% accuracy with 88% non-TB recall. That means we're strong at clearing low-risk cases and reducing false alarms.

Our sputum model — ResNet18 AFB classifier — reached 96.19% AFB-positive sensitivity. It's designed to miss fewer true positives.

System testing: full workflow pass on iPhone 15 and Redmi Note 12 Pro 5G. Quality gates reject fake coughs, silence, and bad smear photos.

Real users: 6 evaluators — nurses, nursing students, an IT specialist, a clinical instructor, and an MLS student. ISO/IEC 25010 grand mean: 4.53 / 5.00 — Strongly Agree. Overall ratings: 33% Excellent, 50% Good, 17% Fair. Usability 4.57. Security 4.94.


[Slide 17 — The honest part]

We're not claiming we cured TB. Here's what we're not: not a diagnostic tool — GeneXpert and smear microscopy still confirm; not replacing clinicians — health workers review every result; not perfect in rural Wi-Fi — that's our #1 field feedback.

What we are: a triage filter that helps overloaded staff prioritize who goes to the lab today; a system that works where labs don't; built on public datasets and real UAT from Filipino health workers.

One evaluator rated us Fair. They told us exactly what we need: better onboarding, location-based DOTS referral, and offline mode. We listened. That's in the roadmap.


[Slide 18 — Future Improvement]

Future improvement on three tracks. Connectivity and deployment: handle unstable rural Wi-Fi, offline queue with sync, faster cloud inference with retry feedback. User experience: onboarding tutorial, location-based DOTS referral, UI polish from evaluator feedback. Models and data: expand cough TB recall, collect more AFB-positive sputum samples, retrain on multi-site clinical data.


[Slide 19 — Evaluation & Scale]

Our evaluation and scale plan: complete additional med-tech UAT and broaden the pool to RHUs, DOTS nurses, and med techs; conduct a field pilot in barangay and RHU settings with real screening volume; long term, integrate with national TB program workflows for referral tracking and reporting.

TBhon remains a pre-screening and triage tool — aligned with GeneXpert, smear microscopy, and X-ray referral pathways. We support the lab. We don't replace it.


[Slide 20 — System Demo + Close]

TB doesn't wait for the lab. Neither should we.

Every day someone in Mindanao coughs into a handkerchief and hopes it's just a cold. TBhon gives health workers the confidence to say: You need to go — today.

[Member 2 name] will now play our system demo — login, checklist, cough capture, sputum, and the fused result.

[After demo — Member 1 closes:]

We're TBhon. Screen smarter. Refer faster. Save lives. Thank you.\
"""


def shade_cell(cell, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_run(run, *, bold=False, size=11, color=None, italic=False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_para(doc: Document, text: str = "", *, bold=False, size=11, align=None, italic=False) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run(run, bold=bold, size=size, italic=italic)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D6EAF8") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, bold=True, size=10)
        shade_cell(cell, header_fill)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run(run, size=10)
    doc.add_paragraph()


def add_script_block(doc: Document, text: str, size: int = 12) -> None:
    """Render script as normal paragraphs — not bullets."""
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("[Slide") or block.startswith("[After"):
            p = doc.add_paragraph()
            run = p.add_run(block)
            set_run(run, bold=True, size=11, color=RGBColor(91, 79, 207))
            continue
        p = doc.add_paragraph()
        run = p.add_run(block.replace("\n", " "))
        set_run(run, size=size)


def add_member_section(
    doc: Document,
    label: str,
    name_slot: str,
    slides: str,
    duration: str,
    script: str,
    handoff: str | None = None,
) -> None:
    doc.add_heading(label, level=1)
    add_para(doc, f"Name: {name_slot}    |    Slides: {slides}    |    ~{duration}", bold=True)
    doc.add_heading("Full script — read this aloud", level=2)
    add_script_block(doc, script)
    if handoff:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(handoff)
        set_run(run, bold=True, size=11, italic=True, color=RGBColor(91, 79, 207))
    doc.add_page_break()


def build_doc() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_heading("TBhon — Pitch Deck Speaking Script", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "3 members  |  20 slides  |  Read the script below — not bullet notes",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
    )
    add_para(doc, "")

    add_table(
        doc,
        ["Member", "Name (fill in)", "Slides", "Time"],
        [
            ["Member 1", "_________________________", "1–6", "~5 min"],
            ["Member 2", "_________________________", "7–15", "~8 min"],
            ["Member 3", "_________________________", "16–20", "~5 min + demo"],
        ],
        header_fill="EDE7F6",
    )

    add_para(doc, "Before you present", bold=True, size=12)
    add_para(
        doc,
        "Replace [Member 2 name] and [Member 3 name] in the scripts with your teammates' names. "
        "Slide 4 is a duplicate title — one line is enough. Slide 14: say the three real risk bands "
        "(Low <38%, Moderate 38–62%, High ≥62%) — the graphic on the slide is outdated. "
        "Slide 20: play the embedded YouTube demo; have a backup if Wi-Fi fails.",
        size=11,
    )

    doc.add_page_break()

    add_member_section(
        doc,
        "MEMBER 1 — Problem & Product Intro",
        "_________________________",
        "1–6",
        "5 min",
        MEMBER_1_SCRIPT,
        handoff="→ Hand off to Member 2 at end of Slide 6.",
    )

    add_member_section(
        doc,
        "MEMBER 2 — Tech, Methods, Fusion, IoT",
        "_________________________",
        "7–15",
        "8 min",
        MEMBER_2_SCRIPT,
        handoff="→ Hand off to Member 3 at end of Slide 15.",
    )

    add_member_section(
        doc,
        "MEMBER 3 — Validation, Honest, Future, Demo, Close",
        "_________________________",
        "16–20",
        "5 min + demo",
        MEMBER_3_SCRIPT,
        handoff="→ Member 2 plays demo on Slide 20. Member 1 delivers final close after demo.",
    )

    doc.add_heading("Slide 20 — Demo only (if not using YouTube)", level=2)
    add_script_block(
        doc,
        """If you run live instead of the video, Member 2 says:

"I'll walk through a live screening now — staff login, patient registration, the 11-item checklist, 3 cough recordings with quality check, sputum capture, and the fused Low, Moderate, or High result on screen."

Then tap through the app. Keep it under 5 minutes.""",
        size=12,
    )

    return doc


def main() -> int:
    doc = build_doc()
    try:
        doc.save(str(OUT))
        print(f"Saved: {OUT}")
    except PermissionError:
        doc.save(str(OUT_ALT))
        print(f"Saved (fallback): {OUT_ALT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
