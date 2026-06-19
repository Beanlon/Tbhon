"""Build Appendix B from real dataset files (sputum images + cough audio).

Copies selected samples out of the project datasets into docs/appendix_b/ and
assembles a Word appendix. Does not synthesize or mock dataset content.
"""
from __future__ import annotations

import csv
import json
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw
from scipy.io import wavfile

ROOT = Path(__file__).resolve().parents[1]
SPUTUM_DATASET = ROOT / "ml (phlegm)" / "Raw_Sputum_Microscopy_Dataset"
APPENDIX_DIR = ROOT / "docs" / "appendix_b"
OUT_DOCX = ROOT / "docs" / "TBhon_Appendix_B.docx"
OUT_DOCX_FALLBACK = ROOT / "docs" / "TBhon_Appendix_B_generated.docx"
OUT_MANIFEST = ROOT / "docs" / "appendix_b_samples.json"

COUGH_DATASET_SLUG = "ruchikashirsath/tb-audio"
COUGH_FOLD = 0
COUGH_SAMPLE_RATE = 16000
COUGH_CLIP_SECONDS = 6.0

IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff"}
GRID_COLS = 5
GRID_ROWS = 4
SAMPLE_COUNT = GRID_COLS * GRID_ROWS


@dataclass(frozen=True)
class SputumSample:
    index: int
    source_image: Path
    source_label: Path
    afb_count: int
    exported_image: Path
    exported_label: Path | None


@dataclass(frozen=True)
class CoughSample:
    index: int
    tb_status: int
    participant: str
    source_wav: Path
    exported_wav: Path
    exported_waveform: Path


def count_boxes(label_path: Path) -> int:
    if not label_path.exists():
        return 0
    count = 0
    with label_path.open("r", encoding="utf-8", errors="ignore") as fh:
        for line in fh:
            parts = line.strip().split()
            if len(parts) >= 5:
                try:
                    float(parts[1])
                    count += 1
                except ValueError:
                    continue
    return count


def parse_yolo_boxes(label_path: Path) -> list[tuple[float, float, float, float]]:
    boxes: list[tuple[float, float, float, float]] = []
    if not label_path.exists():
        return boxes
    with label_path.open("r", encoding="utf-8", errors="ignore") as fh:
        for line in fh:
            parts = line.strip().split()
            if len(parts) < 5:
                continue
            try:
                _, xc, yc, w, h = parts[:5]
                boxes.append((float(xc), float(yc), float(w), float(h)))
            except ValueError:
                continue
    return boxes


def draw_annotated_image(img_path: Path, label_path: Path) -> Image.Image:
    img = Image.open(img_path).convert("RGB")
    w, h = img.size
    draw = ImageDraw.Draw(img)
    for xc, yc, bw, bh in parse_yolo_boxes(label_path):
        x1 = (xc - bw / 2) * w
        y1 = (yc - bh / 2) * h
        x2 = (xc + bw / 2) * w
        y2 = (yc + bh / 2) * h
        draw.rectangle([x1, y1, x2, y2], outline="#E53935", width=max(2, int(min(w, h) * 0.004)))
    return img


def is_usable_image(img_path: Path) -> bool:
    arr = np.asarray(Image.open(img_path).convert("RGB"))
    if arr.size == 0:
        return False
    return float(arr.mean()) > 8.0 and float(arr.std()) > 4.0


def pick_evenly(items: list, n: int) -> list:
    if len(items) <= n:
        return items
    if n <= 1:
        return items[:n]
    indices = [round(i * (len(items) - 1) / (n - 1)) for i in range(n)]
    seen: set[int] = set()
    chosen: list = []
    for idx in indices:
        if idx not in seen:
            seen.add(idx)
            chosen.append(items[idx])
    if len(chosen) < n:
        for item in items:
            if item not in chosen:
                chosen.append(item)
            if len(chosen) == n:
                break
    return chosen


def iter_sputum_pairs(split: str) -> list[tuple[Path, Path, int]]:
    images_dir = SPUTUM_DATASET / "images" / split
    labels_dir = SPUTUM_DATASET / "labels" / split
    pairs: list[tuple[Path, Path, int]] = []
    if not images_dir.is_dir():
        return pairs
    for img in sorted(images_dir.iterdir()):
        if img.suffix.lower() not in IMAGE_SUFFIXES:
            continue
        lbl = labels_dir / f"{img.stem}.txt"
        pairs.append((img, lbl, count_boxes(lbl)))
    return pairs


def resolve_cough_dataset_root() -> Path:
    cache_guess = (
        Path.home()
        / ".cache"
        / "kagglehub"
        / "datasets"
        / "ruchikashirsath"
        / "tb-audio"
        / "versions"
        / "1"
        / "Tuberculosis"
    )
    if cache_guess.is_dir() and (cache_guess / "raw_data").is_dir():
        return cache_guess

    try:
        import kagglehub

        downloaded = Path(kagglehub.dataset_download(COUGH_DATASET_SLUG))
        tb_root = downloaded / "Tuberculosis"
        if tb_root.is_dir():
            return tb_root
    except Exception as exc:  # pragma: no cover - optional runtime dependency
        raise FileNotFoundError(
            "Cough dataset not found. Download ruchikashirsath/tb-audio via kagglehub first."
        ) from exc

    raise FileNotFoundError(f"Expected Tuberculosis dataset root under {downloaded}")


def read_cough_split_csv(path: Path) -> list[tuple[str, int, str]]:
    rows: list[tuple[str, int, str]] = []
    with path.open(newline="", encoding="utf-8", errors="ignore") as fh:
        reader = csv.DictReader(fh)
        for row in reader:
            filename = (row.get("filename") or "").strip()
            participant = (row.get("participant") or "").strip()
            tb_status = int((row.get("tb_status") or "0").strip())
            if filename:
                rows.append((filename, tb_status, participant))
    return rows


def index_cough_wavs(raw_data_dir: Path) -> dict[str, Path]:
    index: dict[str, Path] = {}
    for wav in raw_data_dir.rglob("*.wav"):
        if wav.name not in index:
            index[wav.name] = wav
    if not index:
        raise FileNotFoundError(f"No .wav files found under {raw_data_dir}")
    return index


def collect_sputum_candidates(positive: bool) -> list[tuple[Path, Path, int]]:
    items: list[tuple[Path, Path, int]] = []
    for split in ("train", "val", "test"):
        for img, lbl, count in iter_sputum_pairs(split):
            if not is_usable_image(img):
                continue
            if positive and count > 0:
                items.append((img, lbl, count))
            if not positive and count == 0:
                items.append((img, lbl, count))
    if positive:
        items.sort(key=lambda t: t[2])
    return pick_evenly(items, SAMPLE_COUNT)


def collect_cough_candidates(tb_status: int) -> list[tuple[Path, int, str]]:
    tb_root = resolve_cough_dataset_root()
    meta_csv = tb_root / "metadata" / f"X_train_Fold_{COUGH_FOLD}.csv"
    if not meta_csv.is_file():
        raise FileNotFoundError(f"Missing cough split CSV: {meta_csv}")

    wav_index = index_cough_wavs(tb_root / "raw_data")
    matched: list[tuple[Path, int, str]] = []
    for filename, status, participant in read_cough_split_csv(meta_csv):
        if status != tb_status:
            continue
        wav_path = wav_index.get(filename)
        if wav_path is not None:
            matched.append((wav_path, status, participant))

    matched.sort(key=lambda t: t[0].name)
    return pick_evenly(matched, SAMPLE_COUNT)


def reset_appendix_dir() -> None:
    if APPENDIX_DIR.exists():
        shutil.rmtree(APPENDIX_DIR)
    for sub in (
        "sputum/no_afb/images",
        "sputum/no_afb/labels",
        "sputum/afb_positive/images",
        "sputum/afb_positive/labels",
        "cough/no_tb/wav",
        "cough/no_tb/waveforms",
        "cough/tb/wav",
        "cough/tb/waveforms",
    ):
        (APPENDIX_DIR / sub).mkdir(parents=True, exist_ok=True)


def export_sputum_samples(
    negatives: list[tuple[Path, Path, int]],
    positives: list[tuple[Path, Path, int]],
) -> tuple[list[SputumSample], list[SputumSample]]:
    exported_neg: list[SputumSample] = []
    exported_pos: list[SputumSample] = []

    for i, (src_img, src_lbl, count) in enumerate(negatives, start=1):
        dst_img = APPENDIX_DIR / "sputum/no_afb/images" / f"no_afb_sample_{i:02d}{src_img.suffix.lower()}"
        dst_lbl = APPENDIX_DIR / "sputum/no_afb/labels" / f"no_afb_sample_{i:02d}.txt"
        shutil.copy2(src_img, dst_img)
        if src_lbl.is_file():
            shutil.copy2(src_lbl, dst_lbl)
        exported_neg.append(
            SputumSample(i, src_img, src_lbl, count, dst_img, dst_lbl if src_lbl.is_file() else None)
        )

    for i, (src_img, src_lbl, count) in enumerate(positives, start=1):
        dst_img = APPENDIX_DIR / "sputum/afb_positive/images" / f"afb_positive_sample_{i:02d}.png"
        dst_lbl = APPENDIX_DIR / "sputum/afb_positive/labels" / f"afb_positive_sample_{i:02d}.txt"
        annotated = draw_annotated_image(src_img, src_lbl)
        annotated.save(dst_img)
        if src_lbl.is_file():
            shutil.copy2(src_lbl, dst_lbl)
        exported_pos.append(
            SputumSample(i, src_img, src_lbl, count, dst_img, dst_lbl if src_lbl.is_file() else None)
        )

    return exported_neg, exported_pos


def load_wav_mono(path: Path) -> tuple[np.ndarray, int]:
    sr, data = wavfile.read(path)
    arr = np.asarray(data)
    if arr.ndim == 2:
        arr = arr.mean(axis=1)
    if np.issubdtype(arr.dtype, np.integer):
        max_val = float(np.iinfo(arr.dtype).max)
        arr = arr.astype(np.float32) / max_val
    else:
        arr = arr.astype(np.float32)
    return arr, int(sr)


def resample_if_needed(wav: np.ndarray, sr: int, target_sr: int) -> np.ndarray:
    if sr == target_sr:
        return wav
    try:
        import torch
        import torchaudio

        t = torch.from_numpy(wav).float().unsqueeze(0)
        t = torchaudio.functional.resample(t, sr, target_sr)
        return t.squeeze(0).numpy()
    except Exception:
        duration = len(wav) / sr
        target_len = max(1, int(round(duration * target_sr)))
        x_old = np.linspace(0.0, 1.0, num=len(wav), endpoint=False)
        x_new = np.linspace(0.0, 1.0, num=target_len, endpoint=False)
        return np.interp(x_new, x_old, wav).astype(np.float32)


def clip_or_pad(wav: np.ndarray, target_len: int) -> np.ndarray:
    if len(wav) >= target_len:
        start = max(0, (len(wav) - target_len) // 2)
        return wav[start : start + target_len]
    out = np.zeros(target_len, dtype=np.float32)
    out[: len(wav)] = wav
    return out


def save_waveform_preview(wav_path: Path, out_png: Path) -> None:
    wav, sr = load_wav_mono(wav_path)
    wav = resample_if_needed(wav, sr, COUGH_SAMPLE_RATE)
    target_len = int(COUGH_SAMPLE_RATE * COUGH_CLIP_SECONDS)
    wav = clip_or_pad(wav, target_len)
    times = np.arange(len(wav)) / COUGH_SAMPLE_RATE

    fig, ax = plt.subplots(figsize=(3.2, 1.2), facecolor="white")
    ax.plot(times, wav, color="#1565C0", linewidth=0.7)
    ax.set_xlim(0, COUGH_CLIP_SECONDS)
    ax.set_xlabel("Time (s)", fontsize=7)
    ax.set_ylabel("Amp", fontsize=7)
    ax.tick_params(labelsize=6)
    ax.grid(True, alpha=0.25, linewidth=0.5)
    fig.tight_layout(pad=0.3)
    out_png.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_png, dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def export_cough_samples(
    no_tb: list[tuple[Path, int, str]],
    tb: list[tuple[Path, int, str]],
) -> tuple[list[CoughSample], list[CoughSample]]:
    exported_no_tb: list[CoughSample] = []
    exported_tb: list[CoughSample] = []

    for i, (src_wav, status, participant) in enumerate(no_tb, start=1):
        dst_wav = APPENDIX_DIR / "cough/no_tb/wav" / f"no_tb_sample_{i:02d}.wav"
        dst_png = APPENDIX_DIR / "cough/no_tb/waveforms" / f"no_tb_sample_{i:02d}.png"
        shutil.copy2(src_wav, dst_wav)
        save_waveform_preview(dst_wav, dst_png)
        exported_no_tb.append(CoughSample(i, status, participant, src_wav, dst_wav, dst_png))

    for i, (src_wav, status, participant) in enumerate(tb, start=1):
        dst_wav = APPENDIX_DIR / "cough/tb/wav" / f"tb_sample_{i:02d}.wav"
        dst_png = APPENDIX_DIR / "cough/tb/waveforms" / f"tb_sample_{i:02d}.png"
        shutil.copy2(src_wav, dst_wav)
        save_waveform_preview(dst_wav, dst_png)
        exported_tb.append(CoughSample(i, status, participant, src_wav, dst_wav, dst_png))

    return exported_no_tb, exported_tb


def _style_doc(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def _add_section_heading(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def _add_image_grid(doc: Document, image_paths: list[Path], captions: list[str], thumb_in: float = 1.05) -> None:
    rows = GRID_ROWS
    cols = GRID_COLS
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    for i, (img_path, caption) in enumerate(zip(image_paths, captions)):
        r, c = divmod(i, cols)
        cell = table.rows[r].cells[c]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if img_path.is_file():
            p.add_run().add_picture(str(img_path), width=Inches(thumb_in))
        cap = cell.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(8)
    doc.add_paragraph()


def save_manifest(
    negatives: list[SputumSample],
    positives: list[SputumSample],
    no_tb: list[CoughSample],
    tb: list[CoughSample],
) -> None:
    payload = {
        "sputum_dataset": str(SPUTUM_DATASET),
        "cough_dataset_slug": COUGH_DATASET_SLUG,
        "cough_fold_csv": f"X_train_Fold_{COUGH_FOLD}.csv",
        "negative_sputum_samples": [
            {
                "index": s.index,
                "source_image": str(s.source_image),
                "source_label": str(s.source_label),
                "exported_image": str(s.exported_image),
                "afb_count": s.afb_count,
            }
            for s in negatives
        ],
        "positive_sputum_samples": [
            {
                "index": s.index,
                "source_image": str(s.source_image),
                "source_label": str(s.source_label),
                "exported_image": str(s.exported_image),
                "afb_count": s.afb_count,
            }
            for s in positives
        ],
        "no_tb_cough_samples": [
            {
                "index": s.index,
                "participant": s.participant,
                "source_wav": str(s.source_wav),
                "exported_wav": str(s.exported_wav),
                "exported_waveform": str(s.exported_waveform),
            }
            for s in no_tb
        ],
        "tb_cough_samples": [
            {
                "index": s.index,
                "participant": s.participant,
                "source_wav": str(s.source_wav),
                "exported_wav": str(s.exported_wav),
                "exported_waveform": str(s.exported_waveform),
            }
            for s in tb
        ],
    }
    OUT_MANIFEST.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def save_word_appendix(
    negatives: list[SputumSample],
    positives: list[SputumSample],
    no_tb: list[CoughSample],
    tb: list[CoughSample],
) -> Path:
    doc = Document()
    _style_doc(doc)

    title = doc.add_heading("APPENDIX B — DATASET SAMPLES", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_body(
        doc,
        "This appendix lists representative samples copied directly from the TBhon training datasets: "
        "the local Raw_Sputum_Microscopy_Dataset (YOLO AFB detection) and the Kaggle CODA TB cough "
        "audio corpus (ruchikashirsath/tb-audio). Exported files are stored under docs/appendix_b/.",
    )

    _add_section_heading(doc, "B.1  No AFB Sputum Smear Samples")
    _add_body(
        doc,
        "Representative microscopy fields with zero annotated AFB boxes, taken from train/val/test splits "
        "of Raw_Sputum_Microscopy_Dataset.",
    )
    _add_image_grid(
        doc,
        [s.exported_image for s in negatives],
        [f"No AFB Sample {s.index}" for s in negatives],
    )
    _add_body(
        doc,
        "Note: Images are copied verbatim from the dataset. Matching YOLO label files (empty) are in "
        "docs/appendix_b/sputum/no_afb/labels/.",
    )

    _add_section_heading(doc, "B.2  AFB-Positive Sputum Smear Samples")
    _add_body(
        doc,
        "Representative AFB-positive fields with YOLO bounding boxes drawn from each image's label file. "
        "Original images and .txt annotations are preserved under docs/appendix_b/sputum/afb_positive/.",
    )
    _add_image_grid(
        doc,
        [s.exported_image for s in positives],
        [f"AFB-Positive Sample {s.index} ({s.afb_count} AFB)" for s in positives],
    )
    _add_body(
        doc,
        "Legend: Red boxes = annotated AFB regions (YOLO class 0).",
    )

    _add_section_heading(doc, "B.3  No-TB Cough Audio Samples")
    _add_body(
        doc,
        "Representative cough recordings labeled tb_status=0 from X_train_Fold_0.csv. Each sample includes "
        "the copied .wav file and a waveform preview of the first 6 seconds at 16 kHz.",
    )
    _add_image_grid(
        doc,
        [s.exported_waveform for s in no_tb],
        [f"No-TB Cough Sample {s.index}\n{s.source_wav.name}" for s in no_tb],
        thumb_in=1.15,
    )
    _add_body(
        doc,
        "Note: Audio files are copied verbatim to docs/appendix_b/cough/no_tb/wav/.",
    )

    _add_section_heading(doc, "B.4  TB-Positive Cough Audio Samples")
    _add_body(
        doc,
        "Representative cough recordings labeled tb_status=1 from X_train_Fold_0.csv. Each sample includes "
        "the copied .wav file and a waveform preview of the first 6 seconds at 16 kHz.",
    )
    _add_image_grid(
        doc,
        [s.exported_waveform for s in tb],
        [f"TB Cough Sample {s.index}\n{s.source_wav.name}" for s in tb],
        thumb_in=1.15,
    )
    _add_body(
        doc,
        "Note: Audio files are copied verbatim to docs/appendix_b/cough/tb/wav/.",
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(OUT_DOCX))
        return OUT_DOCX
    except PermissionError:
        doc.save(str(OUT_DOCX_FALLBACK))
        print(f"[warn] Could not overwrite {OUT_DOCX} (file may be open). Saved to {OUT_DOCX_FALLBACK}")
        return OUT_DOCX_FALLBACK


def main() -> int:
    if not SPUTUM_DATASET.is_dir():
        print(f"Sputum dataset not found: {SPUTUM_DATASET}", file=sys.stderr)
        return 1

    try:
        resolve_cough_dataset_root()
    except FileNotFoundError as exc:
        print(str(exc), file=sys.stderr)
        return 1

    negatives_raw = collect_sputum_candidates(positive=False)
    positives_raw = collect_sputum_candidates(positive=True)
    no_tb_raw = collect_cough_candidates(tb_status=0)
    tb_raw = collect_cough_candidates(tb_status=1)

    for label, items in (
        ("negative sputum", negatives_raw),
        ("positive sputum", positives_raw),
        ("no-TB cough", no_tb_raw),
        ("TB cough", tb_raw),
    ):
        if len(items) < SAMPLE_COUNT:
            print(f"Need {SAMPLE_COUNT} {label} samples, found {len(items)}.", file=sys.stderr)
            return 1

    reset_appendix_dir()
    negatives, positives = export_sputum_samples(negatives_raw, positives_raw)
    no_tb, tb = export_cough_samples(no_tb_raw, tb_raw)
    save_manifest(negatives, positives, no_tb, tb)
    docx_path = save_word_appendix(negatives, positives, no_tb, tb)

    print("Exported real dataset samples to:")
    print(f"  {APPENDIX_DIR}")
    print("Generated:")
    print(f"  {docx_path}")
    print(f"  {OUT_MANIFEST}")
    print(f"\nSputum: {len(negatives)} no-AFB, {len(positives)} AFB-positive")
    print(f"Cough:  {len(no_tb)} no-TB, {len(tb)} TB-positive")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
