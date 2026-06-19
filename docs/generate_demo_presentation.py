"""Generate TBhon System Demo Script + Flow (Word) — 3 members, full speaking lines."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "TBhon_Demo_Presentation.docx"
OUT_ALT = ROOT / "docs" / "TBhon_Demo_Presentation_generated.docx"


def shade_cell(cell, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_run(run, *, bold=False, size=11, color=None, italic=False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_para(doc: Document, text: str = "", *, bold=False, size=11, align=None, italic=False) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run(run, bold=bold, size=size, italic=italic)


def add_bullets(doc: Document, items: list[str], *, size=11) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run(run, size=size)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D6EAF8") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, bold=True, size=10)
        shade_cell(cell, header_fill)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run(run, size=10)
    doc.add_paragraph()


def add_banner(doc: Document) -> None:
    table = doc.add_table(rows=2, cols=1)
    banner = table.rows[0].cells[0]
    shade_cell(banner, "0B1530")
    banner.text = ""
    run = banner.paragraphs[0].add_run("TBhon — System Demo Script & Flow")
    set_run(run, bold=True, size=20, color=RGBColor(255, 255, 255))
    banner.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = table.rows[1].cells[0]
    shade_cell(sub, "5B4FCF")
    sub.text = ""
    run = sub.paragraphs[0].add_run("3 Members  |  Live App Walkthrough  |  ~10–12 min")
    set_run(run, bold=True, size=12, color=RGBColor(255, 255, 255))
    sub.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def add_script_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        if not line.strip():
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run(run, size=12)


def add_member_script(
    doc: Document,
    member_label: str,
    name_slot: str,
    duration: str,
    flow_rows: list[list[str]],
    script_lines: list[str],
    handoff: str | None = None,
) -> None:
    doc.add_heading(f"{member_label} — Speaking Script", level=1)
    add_para(doc, f"Name: {name_slot}    |    Duration: {duration}", bold=True)

    doc.add_heading("What to tap (quick reference)", level=2)
    add_table(doc, ["#", "Screen", "Tap / Do", "Say while doing it"], flow_rows)

    doc.add_heading("Full script (read this)", level=2)
    add_script_block(doc, script_lines)

    if handoff:
        p = doc.add_paragraph()
        run = p.add_run(handoff)
        set_run(run, bold=True, size=12, color=RGBColor(91, 79, 207))
    doc.add_page_break()


def build() -> Document:
    doc = Document()
    add_banner(doc)

    doc.add_heading("Team", level=1)
    add_table(
        doc,
        ["Member", "Name (fill in)", "Part"],
        [
            ["Member 1", "_________________________", "Login → intake → checklist"],
            ["Member 2", "_________________________", "Cough → sputum → review → processing"],
            ["Member 3", "_________________________", "Staff review → result → history"],
        ],
        header_fill="E8DAEF",
    )

    doc.add_heading("Demo patient — enter this data", level=2)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Name", "Juan Dela Cruz"],
            ["Birthdate", "15 Mar 1990"],
            ["Gender", "Male"],
            ["Address", "123 Rizal St, Brgy. San Jose, Quezon City"],
            ["Contact", "09171234567"],
            ["Checklist (answer Yes)", "Cough 2+ weeks, night sweats, weight loss, TB contact"],
        ],
    )

    doc.add_heading("Before demo", level=2)
    add_bullets(
        doc,
        [
            "Staff logged in · ESP32 on + same Wi-Fi as phone",
            "Sputum smear ready on slide",
            "Backend + ML online (Processing must finish)",
        ],
    )

    doc.add_page_break()

    # ── MEMBER 1 ──
    add_member_script(
        doc,
        "MEMBER 1",
        "_________________________",
        "~4 min",
        [
            ["1", "Login", "Sign in", '"I log in as booth staff."'],
            ["2", "Home", "Start screening", '"I start a new booth session."'],
            ["3", "Staff instructions", "Start session", '"Phone and booth device are on the same Wi-Fi."'],
            ["4", "Patient type", "New patient", '"This is a first-time walk-in."'],
            ["5", "Client intake", "Fill identity, address, contact → Review → Continue", '"I enter the patient record."'],
            ["6", "Checklist", "Answer 11 questions", '"Staff asks each symptom and risk question."'],
            ["7", "Checklist summary", "Continue to device cough capture", '"Checklist done — next is cough on the booth device."'],
        ],
        [
            "I log in as booth staff and open the home screen.",
            "",
            'I tap Start screening. The staff instructions remind us that this phone runs the session and the booth device captures cough audio.',
            "",
            'I tap Start session. The app creates a new screening record on the server.',
            "",
            'On Patient type, I choose New patient — first time at this booth.',
            "",
            "On Client intake I enter the patient's details:",
            "— Name: Juan Dela Cruz",
            "— Birthdate: March 15, 1990",
            "— Gender: Male",
            "— Address: 123 Rizal St, Brgy. San Jose, Quezon City",
            "— Contact: 09171234567",
            "",
            "I skip emergency contact and government ID for this demo, review the summary, and tap Continue.",
            "",
            "Next is the symptom checklist — eleven yes-or-no questions, one per screen. I ask the patient and tap their answer.",
            "",
            "For this demo we mark Yes on: cough lasting two weeks or longer, night sweats, unexplained weight loss, and close contact with someone who has TB.",
            "",
            "On the checklist summary the app shows how many symptoms were reported and a concern level.",
            "",
            'I tap Continue to device cough capture.',
            "",
            "That's my part — [Name 2] will run cough and sputum capture on the booth device.",
        ],
        handoff="► HANDOFF to Member 2 — pass the phone. Booth device ready for 3 coughs.",
    )

    # ── MEMBER 2 ──
    add_member_script(
        doc,
        "MEMBER 2",
        "_________________________",
        "~4 min",
        [
            ["8", "IoT cough", "Patient coughs ×3 on booth device", '"One cough at a time — wait for each slot."'],
            ["9", "IoT cough", "Continue", '"All three cough clips are saved."'],
            ["10", "IoT sputum", "Prepare smear → Start capture", '"Staff prepares the slide, device takes the photo."'],
            ["11", "Review", "Play back coughs, check sputum image", '"We verify everything before analysis."'],
            ["12", "Review", "Tap Analyze", '"Submitting to the server for ML."'],
            ["13", "Processing", "Wait", '"Uploading audio and image — computing fused risk."'],
        ],
        [
            "I'm on the cough capture screen. The patient will cough into the booth device three separate times.",
            "",
            "Slot one… [patient coughs] …saved.",
            "Slot two… [patient coughs] …saved.",
            "Slot three… [patient coughs] …saved.",
            "",
            "All three slots are complete. I tap Continue.",
            "",
            "On sputum capture, I prepare the smear on the slide, then tap Start capture on the booth device.",
            "",
            "[Wait for device photo] The smear image is captured and attached to this session.",
            "",
            "(If sputum is not available: tap Skip, select a reason — fusion will use checklist and cough only.)",
            "",
            "On Review I expand the cough clips and play one back to confirm audio was captured. I also check the sputum thumbnail.",
            "",
            "Everything looks complete. I tap Analyze.",
            "",
            "The Processing screen uploads the cough clips and sputum image, runs cough ML and sputum ML, then fuses them with the checklist into one triage risk.",
            "",
            "[While waiting] The cough model scores each valid clip. The sputum model classifies the smear. The app combines all three into a Low, Moderate, or High risk band.",
            "",
            "Processing is done — Staff review is next.",
            "",
            "[Name 3] will confirm the result and show it to the patient.",
        ],
        handoff="► HANDOFF to Member 3 — Staff review screen is showing.",
    )

    # ── MEMBER 3 ──
    add_member_script(
        doc,
        "MEMBER 3",
        "_________________________",
        "~3 min",
        [
            ["14", "Staff review", "Read triage level + fused %", '"Staff checks the output before the patient sees it."'],
            ["15", "Staff review", "Confirm & show result", '"Optional staff notes, then release result."'],
            ["16", "Result", "Scroll through gauge and breakdown", '"Checklist, cough ML, sputum ML, fused risk."'],
            ["17", "Result", "Show QR on result slip", '"Patient scans this to claim their result later."'],
            ["18", "Result / History", "Session details → booth history", '"Session is saved — demo complete."'],
        ],
        [
            "On Staff review I see the fused triage level and TB probability before the patient views anything.",
            "",
            "For this session it shows [read the level on screen — Low / Moderate / High] at [read the percentage]% fused probability.",
            "",
            "If the risk is Moderate or High, the referral line appears — refer for GeneXpert, smear, or clinical workup.",
            "",
            "I can add optional staff notes here. Then I tap Confirm and show result.",
            "",
            "This is the result screen. At the top is the fused risk gauge.",
            "",
            "Scrolling down: checklist contribution, cough ML scores from the three clips, sputum ML classification, and the fusion breakdown showing how each modality contributed.",
            "",
            "The disclaimer states this is pre-screening support — not a diagnosis.",
            "",
            "Here is the patient QR on the result slip. The patient scans this on their own phone later to open their screening result.",
            "",
            "I can also export a PDF summary if the booth needs a printed copy.",
            "",
            "Opening session details — full record with cough playback, checklist answers, and timestamps.",
            "",
            "Back on Home, booth session history lists this visit. The screening is complete.",
            "",
            "Thank you — we're open for questions.",
        ],
        handoff=None,
    )

    doc.add_heading("Backup if something fails", level=1)
    add_table(
        doc,
        ["Problem", "Say this + do this"],
        [
            [
                "Processing hangs",
                '"While that loads, I\'ll show a completed session from history." → open History',
            ],
            [
                "Cough slot fails",
                '"Let me re-record that slot." → redo one cough on device',
            ],
            [
                "Sputum fails",
                '"We\'ll skip sputum for this run." → Skip with reason',
            ],
            [
                "ESP32 offline",
                '"The booth device isn\'t connected — I\'ll walk through a finished session instead." → History',
            ],
        ],
        header_fill="FADBD8",
    )

    add_para(doc, "TBhon v1.2 — rehearse with real device once before panel.", italic=True, size=10)
    return doc


def main() -> None:
    doc = build()
    for path in (OUT, OUT_ALT):
        try:
            doc.save(path)
            print(f"Wrote {path}")
            return
        except PermissionError:
            continue
    doc.save(OUT_ALT)
    print(f"Wrote {OUT_ALT} only (close open Word files and re-run to update main file)")


if __name__ == "__main__":
    main()
