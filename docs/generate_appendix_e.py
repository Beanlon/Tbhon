"""Generate TBhon Appendix E — Model Training Repository (Word)."""
from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
APPENDIX_E = ROOT / "docs" / "appendix_e"
FIG = ROOT / "docs" / "figures"
OUT = ROOT / "docs" / "TBhon_Appendix_E.docx"
OUT_FALLBACK = ROOT / "docs" / "TBhon_Appendix_E_generated.docx"
FIG_SCRIPT = ROOT / "docs" / "generate_appendix_e_figures.py"

COUGH_CONFIG = ROOT / "ml" / "runs" / "20260531_014419" / "config.json"
SPUTUM_CONFIG = ROOT / "ml (phlegm)" / "runs" / "phlegm_afb_binary_20260531_133949" / "config.json"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_code_block(doc: Document, text: str) -> None:
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def add_figure(doc: Document, path: Path, width: float = 6.0) -> None:
    if path.is_file():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"[Insert figure: {path.name}]")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    doc.add_paragraph()


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_figures() -> None:
    if not (APPENDIX_E / "e3_repository_structure.png").is_file():
        subprocess.run([sys.executable, str(FIG_SCRIPT)], check=True, cwd=str(ROOT))


def build() -> None:
    ensure_figures()
    cough = load_json(COUGH_CONFIG)
    sputum = load_json(SPUTUM_CONFIG)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    add_heading(doc, "APPENDIX E", 1)
    add_heading(doc, "MODEL TRAINING REPOSITORY", 1)
    add_para(
        doc,
        "This appendix documents the model training repository, configuration parameters, and training "
        "environment used to produce the production cough hybrid classifier (run 20260531_014419) and "
        "sputum AFB binary ResNet18 classifier (run phlegm_afb_binary_20260531_133949) reported in Chapter VI.",
    )

    add_heading(doc, "E.1 Repository Structure", 2)
    add_para(
        doc,
        "The TBhon machine learning assets are organized under the main project repository with separate "
        "directories for cough audio and sputum microscopy pipelines. Production checkpoints, metrics, and "
        "configuration files are stored under timestamped run folders.",
    )
    add_code_block(
        doc,
        """Tbhon/
├── ml/
│   ├── train_tb_cough_hybrid.py
│   ├── train_tb_cough_cnn.py
│   ├── infer_api.py
│   ├── production_model.json
│   ├── runs/
│   │   └── 20260531_014419/
│   │       ├── model.pt
│   │       ├── hybrid_bundle.pkl
│   │       ├── metrics.json
│   │       └── config.json
│   └── scripts/
└── ml (phlegm)/
    ├── train_phlegm_cnn.py
    ├── Raw_Sputum_Microscopy_Dataset/
    └── runs/
        └── phlegm_afb_binary_20260531_133949/
            ├── model_best.pt
            ├── metrics.json
            └── config.json""",
    )
    add_figure(doc, APPENDIX_E / "e3_repository_structure.png", 6.2)
    add_caption(doc, "Figure E.1. TBhon Model Training Repository Structure (VS Code Explorer)")

    add_heading(doc, "E.2 Training Configuration", 2)
    add_para(doc, "Table E.2.1. Cough Hybrid Classifier — Production Training Configuration (run 20260531_014419)")
    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Framework", "PyTorch CNN + scikit-learn Gradient Boosting"],
            ["Language", "Python 3"],
            ["Dataset", cough.get("dataset_slug", "ruchikashirsath/tb-audio")],
            ["Cross-validation fold", str(cough.get("fold", 1))],
            ["Sample rate", f"{cough.get('sample_rate', 16000)} Hz"],
            ["Clip duration", f"{cough.get('clip_seconds', 4.0)} s"],
            ["Mel bins", str(cough.get("n_mels", 64))],
            ["CNN epochs", str(cough.get("cnn_epochs", 8))],
            ["CNN batch size", str(cough.get("cnn_batch_size", 32))],
            ["CNN learning rate", str(cough.get("cnn_lr", 0.0003))],
            ["GBM estimators", str(cough.get("gbm_n_estimators", 300))],
            ["Validation fraction", str(cough.get("val_fraction", 0.12))],
            ["Device", "CPU — local Windows workstation (Visual Studio Code)"],
        ],
    )

    add_para(doc, "Table E.2.2. Sputum ResNet18 AFB Binary Classifier — Production Training Configuration")
    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Framework", "PyTorch / torchvision ResNet18"],
            ["Language", "Python 3"],
            ["Task", sputum.get("task", "binary")],
            ["Image size", f"{sputum.get('img_size', 224)}×{sputum.get('img_size', 224)}"],
            ["Epochs", str(sputum.get("epochs", 30))],
            ["Batch size", str(sputum.get("batch_size", 32))],
            ["Learning rate", str(sputum.get("lr", 0.001))],
            ["Weight decay", str(sputum.get("weight_decay", 0.0001))],
            ["Augmentation", str(sputum.get("augment", True))],
            ["Class weighting", str(sputum.get("class_weight", True))],
            ["Optimizer", "Adam"],
            ["Device", str(sputum.get("device", "cpu"))],
        ],
    )

    add_heading(doc, "E.3 Training Environment", 2)
    add_para(
        doc,
        "All production model training documented in Chapter VI was performed on a local Windows 11 development "
        "workstation using Visual Studio Code and PowerShell. The TBhon repository was cloned from GitHub; "
        "Python dependencies were installed from ml/requirements.txt and ml (phlegm) requirements. Cough hybrid "
        "training and sputum ResNet18 training were executed as command-line scripts—no Google Colab or cloud "
        "notebook environment was used for the reported production runs.",
    )

    add_para(
        doc,
        "Figure E.3.1 shows the ML repository structure in the VS Code explorer. Figure E.3.2 shows the local "
        "VS Code terminal session used to run the production cough and sputum training commands.",
    )

    add_figure(doc, APPENDIX_E / "e3_repository_structure.png", 6.2)
    add_caption(doc, "Figure E.3.1. Repository Screenshot — TBhon ML Directory (VS Code Explorer)")

    add_figure(doc, APPENDIX_E / "e3_vscode_training_session.png", 6.2)
    add_caption(doc, "Figure E.3.2. Local Training Environment — Visual Studio Code Terminal (Windows)")

    add_para(
        doc,
        "Training commands (reference):\n"
        "• Cough: python train_tb_cough_hybrid.py --fold 1 --cnn-epochs 8 --gbm-estimators 300\n"
        "• Sputum: python train_phlegm_cnn.py --task binary --epochs 30",
    )

    try:
        doc.save(OUT)
        print(f"Wrote {OUT}")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"Wrote {OUT_FALLBACK} (close {OUT.name} in Word to overwrite primary file)")


if __name__ == "__main__":
    build()
