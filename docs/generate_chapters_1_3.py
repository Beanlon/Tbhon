"""Generate TBhon Chapters 1-3 (HCI/UCSD format) as Word document."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\Mika\Documents\GitHub\Tbhon\docs\TBhon_Chapters_1-3.docx"


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(12)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(12)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(12)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # --- CHAPTER I ---
    add_heading(doc, "CHAPTER I. INTRODUCTION", 1)

    add_heading(doc, "Background of the Study", 2)
    add_para(
        doc,
        "Tuberculosis (TB) is a contagious bacterial infection caused by Mycobacterium tuberculosis that primarily affects the lungs. "
        "It spreads through airborne droplets when an infected person coughs or sneezes. Early symptoms such as persistent cough, "
        "chest pain, and sputum production are often ignored or misdiagnosed as a common respiratory illness. According to the WHO "
        "Global TB Report 2025, the disease claimed over 1.2 million lives and affected an estimated 10.7 million people worldwide, "
        "with 87% of cases concentrated in just 30 countries—including the Philippines, India, Indonesia, and China. Despite advances "
        "in diagnosis and treatment, millions remain undiagnosed because of limited access, cost, stigma, and delayed health-seeking behavior."
    )
    add_para(
        doc,
        "On a global scale, digital health and machine learning have introduced new approaches to early TB screening through non-invasive "
        "modalities such as cough sound analysis and computer-aided sputum microscopy. Studies demonstrate that automated screening can "
        "identify microbiologically confirmed TB cases that would have been missed by symptom screening alone. However, most existing "
        "systems remain single-modality, infrastructure-heavy, or clinically dependent, limiting their applicability in community-based and "
        "resource-constrained settings where early detection is most critical."
    )
    add_para(
        doc,
        "In the Philippines, the Department of Health (DOH) and the World Health Organization (WHO) are intensifying efforts to eliminate "
        "TB under the Philippine Strategic TB Elimination Plan Phase 2 (2025–2030), with a goal of screening 12 million Filipinos by 2026. "
        "The government has increased support for TB programs, including expanded budgets and adoption of AI-powered chest X-rays and rapid "
        "molecular tests. Nevertheless, rural health units (RHUs), barangay health stations, and temporary screening booths still face "
        "constraints in trained personnel, laboratory capacity, and integrated digital workflows for preliminary case prioritization."
    )
    add_para(
        doc,
        "In Davao City and similar urban–periurban settings in Mindanao, community health workers and booth staff must screen walk-in clients "
        "while managing diverse respiratory complaints and limited diagnostic infrastructure. Rapid urbanization, healthcare worker workload, "
        "and fragmented screening tools create an opportunity for an accessible, staff-operated triage support platform that combines "
        "symptom checklist, cough audio analysis, and sputum smear image signals within a single mobile workflow."
    )
    add_para(
        doc,
        "To address these challenges, TBhon was developed as an IoT-enabled, mobile-based multimodal TB pre-screening and triage support system. "
        "It integrates an Expo/React Native staff application, a Node.js/Express backend with Prisma and MySQL, a Python FastAPI machine learning "
        "inference service (PyTorch CNN for cough Mel-spectrograms and sputum microscopy images), optional ESP32-based IoT capture hardware, "
        "and cloud deployment through DigitalOcean and Cloudflare tunnels. The system produces low, moderate, or high triage risk scores through "
        "weighted multimodal fusion and is explicitly designed as triage support—not a medical diagnosis."
    )

    add_heading(doc, "Statement of the Problem", 2)
    add_para(
        doc,
        "Tuberculosis remains a significant global and national health concern, particularly in low- and middle-income communities where "
        "access to screening and diagnostic facilities is limited. Diagnostic delays are influenced by patient-related factors (low symptom "
        "recognition, stigma) and health-system factors (insufficient human resources, limited laboratory capacity, and divided attention "
        "among healthcare workers handling multiple diseases). Preliminary screening support is therefore essential to prioritize cases "
        "requiring further clinical evaluation such as GeneXpert, smear microscopy, or full clinical workup."
    )
    add_para(doc, "Specifically, the following problems were identified:")
    add_bullets(
        doc,
        [
            "Existing TB screening workflows often rely on symptom questionnaires or single-modality tools that lack integrated cough audio and sputum smear analysis in one accessible platform.",
            "Healthcare workers in booth and RHU settings experience fragmented tools—separate paper forms, standalone apps, and laboratory-dependent methods—making consistent triage documentation difficult.",
            "Many AI-assisted TB screening prototypes remain clinically dependent, infrastructure-heavy, or limited to audio-only or imaging-only analysis, reducing real-world deployability in community settings.",
            "There is no unified staff-operated mobile platform in the local context that combines IoT-assisted capture, symptom checklist, multimodal ML fusion, session history, patient result claiming, and referral tracking for TB triage support.",
            "Resource-limited facilities require a low-cost, cloud-connected system that supports staff review, honest scope disclaimers, and traceable screening records without replacing formal diagnosis.",
        ],
    )
    add_para(
        doc,
        "Therefore, this study aimed to develop TBhon, an IoT-enabled mobile-based multimodal diagnostic support system using respiratory "
        "audio and sputum image analysis for tuberculosis pre-screening, integrated with cloud-based data management, staff-operated triage "
        "workflows, and optional ESP32 acquisition hardware to provide accessible, non-invasive, and real-time preliminary TB risk stratification."
    )

    add_heading(doc, "Assumption of the Study", 2)
    add_para(doc, "The following assumptions were made to address the problems identified above:")
    assumptions = [
        (
            "Integration of CNN-based cough audio analysis using Mel-spectrogram representations:",
            "Converting cough recordings into Mel-spectrograms and processing them with convolutional neural networks (CNNs) implemented in PyTorch can identify TB-related acoustic patterns and support preliminary risk assessment when combined with quality gates for recording authenticity.",
        ),
        (
            "Sputum smear image classification using computer vision:",
            "CNN-based analysis of Ziehl–Neelsen stained sputum microscopy images can detect visual TB-related indicators and complement audio-based signals, improving multimodal robustness compared to single-modality screening.",
        ),
        (
            "Multimodal risk fusion through weighted log-odds integration:",
            "Fusing checklist answers, cough ML probabilities, and sputum ML outputs into unified low, moderate, or high triage risk levels will provide staff with actionable prioritization guidance while maintaining honest scope as triage support only.",
        ),
        (
            "Mobile and cloud architecture for booth deployment:",
            "A React Native (Expo) mobile application connected to an Express/Prisma backend and a FastAPI ML inference server—deployed via DigitalOcean and Cloudflare tunnels—can deliver real-time screening support with acceptable response times in field conditions.",
        ),
        (
            "IoT-assisted acquisition via ESP32 modules:",
            "An ESP32-based module for respiratory audio capture and an integrated ESP32–OV5640 microscopic imaging module can support standardized cough and sputum capture in booth settings when phone-only capture is insufficient.",
        ),
        (
            "Staff-operated workflow with patient result access:",
            "Requiring staff review before displaying results, storing sessions in a facility-linked database, and enabling patients to claim results via QR/deep link will improve workflow accountability and follow-up without presenting the tool as self-diagnosis.",
        ),
        (
            "Public dataset training and ISO/IEC-informed evaluation:",
            "Publicly available TB cough audio and sputum smear datasets, combined with standard ML metrics (accuracy, precision, recall, F1-score) and ISO/IEC 25010–aligned software quality assessment, can validate system feasibility prior to formal clinical deployment.",
        ),
    ]
    for title, body in assumptions:
        p = doc.add_paragraph()
        r1 = p.add_run(title + " ")
        r1.bold = True
        r2 = p.add_run(body)
        r1.font.size = r2.font.size = Pt(12)

    add_heading(doc, "Significance of the Study", 2)
    add_para(
        doc,
        "The findings of this study are expected to contribute to the advancement of accessible, multimodal TB pre-screening systems through "
        "the integration of machine learning, mobile computing, IoT acquisition, and cloud-based health informatics. The proposed TBhon "
        "application may provide practical, technological, and societal benefits to the following stakeholders, ranked from highest to "
        "lowest level of expected impact:"
    )
    significance = [
        (
            "1. Community Health Workers, Booth Staff, and Primary Healthcare Units",
            "TBhon may function as a supplementary triage decision-support tool for prioritizing individuals requiring further clinical evaluation. "
            "By combining checklist, cough audio, and sputum smear signals in one staff-operated workflow, the system may improve screening "
            "efficiency in RHUs, barangay health stations, and temporary screening booths with limited personnel and diagnostic resources.",
        ),
        (
            "2. Individuals and At-Risk Populations",
            "The proposed system may encourage earlier health-seeking behavior by surfacing potential TB risk indicators through non-invasive, "
            "accessible screening. Patients may claim their results via QR code, supporting continuity of care and referral to GeneXpert or "
            "clinical workup when triage risk is moderate or high.",
        ),
        (
            "3. Healthcare System and Public Health Services",
            "The study supports scalable digital health solutions that complement existing TB diagnostic procedures. By reducing reliance on "
            "infrastructure-intensive methods during initial screening, the system may contribute to expanding TB screening coverage in "
            "underserved regions aligned with the Philippine Strategic TB Elimination Plan.",
        ),
        (
            "4. Information Technology and Emerging Technology Practitioners",
            "The study may serve as a reference for developers implementing multimodal health AI systems, combining PyTorch inference, FastAPI "
            "services, React Native mobile apps, Prisma/MySQL backends, and ESP32 IoT firmware within a deployable real-world architecture.",
        ),
        (
            "5. Researchers and Future Developers",
            "The study may provide baseline knowledge on multimodal TB screening, dataset utilization, model development, risk fusion, and "
            "system integration—serving as a foundation for future research on lightweight mobile AI, multimodal fusion, and community-level "
            "health screening deployment.",
        ),
        (
            "6. Students and Academic Institutions",
            "The study may contribute to academic learning in Computer Science, Information Technology, and Software Engineering by "
            "demonstrating Human-Computer Interaction (HCI), user-centered design, machine learning, and mobile system development in a "
            "healthcare application context.",
        ),
        (
            "7. Local Government Units and Health Offices in Davao City",
            "TBhon may support local TB elimination initiatives by providing traceable screening sessions, referral tracking, PDF export for "
            "documentation, and in-app TB education content for staff counseling—contributing to technology-assisted public health programs.",
        ),
    ]
    for title, body in significance:
        p = doc.add_paragraph()
        r1 = p.add_run(title + "\n")
        r1.bold = True
        r2 = p.add_run(body)
        r1.font.size = r2.font.size = Pt(12)

    doc.add_page_break()

    # --- CHAPTER II ---
    add_heading(doc, "CHAPTER II", 1)
    add_heading(doc, "RESEARCH DESIGN", 2)
    add_para(
        doc,
        "This chapter presents the User-Centered System Design (UCSD) process model utilized by the researchers in developing TBhon: "
        "An IoT-Enabled Mobile-Based Multimodal Diagnostic Support System Using Respiratory Audio and Sputum Image Analysis for "
        "Tuberculosis Pre-screening. The design process guided the group from initial task analysis through requirements gathering, "
        "prototyping, and evaluation. Each stage reflects both the theoretical foundations of the UCSD model and the researchers' "
        "firsthand experiences during the development process."
    )

    add_heading(doc, "User-Centered System Design Process", 2)
    add_para(
        doc,
        "The User-Centered System Design (UCSD) process is a design philosophy and methodology that prioritizes the needs, limitations, "
        "and preferences of end users at every stage of system development. According to Eason (1988), a UCSD approach requires that "
        "system designers continually focus on users, their tasks, and their working environments throughout the design lifecycle. The "
        "UCSD model followed in this study consists of the following stages: (A) Task Analysis, (B) Requirements Gathering, "
        "(C) Storyboarding and Prototyping, and (D) Evaluation of Prototype."
    )

    add_heading(doc, "A. Task Analysis", 3)
    add_para(
        doc,
        "Task Analysis is the process of understanding and describing the activities users perform to achieve specific goals within a system. "
        "A Hierarchical Task Analysis (HTA) was conducted to decompose the primary goals of TBhon into sub-tasks and actions. This structured "
        "approach helped the research group identify the key interactions booth staff, administrators, and patients must perform when using "
        "the application and served as a blueprint for interface and feature design."
    )
    add_para(
        doc,
        "As the researchers conducted their task analysis, they identified several high-level goals that TBhon must support: "
        "(1) staff-operated TB screening and triage, (2) multimodal data capture (checklist, cough audio, sputum smear), "
        "(3) ML-assisted risk fusion and staff review, (4) session documentation and referral tracking, "
        "(5) patient result claiming, and (6) facility administration. Each of these top-level goals was then broken down into progressively "
        "detailed sub-tasks and operations."
    )
    add_para(
        doc,
        "The researchers found this stage particularly insightful during the early phases of the project. Upon mapping out the screening task, "
        "for example, the group realized that staff must seamlessly transition between client intake, an 11-question symptom/risk checklist, "
        "three cough recordings, optional sputum capture, ML processing, and staff confirmation before showing results—all while maintaining "
        "clear disclaimers that the output is triage support, not diagnosis. This discovery significantly influenced how the group structured "
        "the screening flow of the application."
    )

    add_para(doc, "Table 2.1. Hierarchical Task Analysis of TBhon", bold=True)
    add_table(
        doc,
        ["Level", "Task/Goal", "Sub-tasks", "Operations"],
        [
            ["0", "Use TBhon Application", "—", "Open app, log in as staff, and access home dashboard"],
            [
                "1",
                "Conduct TB Screening Session",
                "1.1 Client intake\n1.2 Complete checklist\n1.3 Capture cough audio\n1.4 Capture sputum (optional)\n1.5 Review and process\n1.6 Staff review and result",
                "Enter client demographics; answer 11 symptom/risk questions; record 3 cough clips (phone or IoT); capture or skip sputum with reason; submit for ML fusion; confirm triage risk before displaying result and QR",
            ],
            [
                "2",
                "Capture Data via IoT Hardware",
                "2.1 Pair device\n2.2 Configure Wi‑Fi\n2.3 Trigger cough capture\n2.4 Trigger sputum capture",
                "Connect ESP32 via BLE; provision network; poll backend for device status; queue capture commands; receive uploads from IoT module",
            ],
            [
                "3",
                "Review and Manage Sessions",
                "3.1 View history\n3.2 Open session details\n3.3 Update referral status\n3.4 Export PDF",
                "Browse past screenings; replay cough audio and view sputum; track referral from recommended to completed; generate triage summary PDF",
            ],
            [
                "4",
                "Claim Screening Result (Patient)",
                "4.1 Scan QR / open link\n4.2 Preview session\n4.3 Claim and view history",
                "Scan patient access QR; preview screening summary; create/link PATIENT account; view claimed results on home screen",
            ],
            [
                "5",
                "Administer Facilities",
                "5.1 Manage facilities\n5.2 Assign staff",
                "Create facility with invite code; assign staff operators; monitor booth-linked screening activity",
            ],
        ],
    )
    add_para(
        doc,
        "Figure 2.1 (see attached prototype figures) presents the full hierarchical task diagram for TBhon, illustrating the flow from "
        "high-level goals down to individual user operations. The HTA reinforced the need for a streamlined, staff-guided interface with "
        "explicit triage disclaimers, which became a guiding principle throughout the design process."
    )

    add_heading(doc, "B. Requirements Gathering", 3)
    add_para(
        doc,
        "Requirements Gathering is the process of collecting the necessary information, user needs, and technical specifications that the "
        "proposed system must fulfill. This stage served as the foundation for all design and development decisions made throughout the project. "
        "The researchers employed three primary methods for gathering requirements: interview, survey/questionnaire, and observation."
    )

    add_heading(doc, "Interview", 4)
    add_para(
        doc,
        "Interviews were conducted with selected individuals including community health workers, nursing students, and booth staff familiar "
        "with TB screening workflows in Davao City. The researchers prepared semi-structured interview guides focused on understanding current "
        "screening practices, frustrations with paper-based or fragmented tools, concerns about false reassurance from automated tools, and "
        "the need for staff-controlled result disclosure."
    )
    add_para(
        doc,
        "During the interview process, a recurring theme was the need for honest scope language—respondents emphasized that any digital tool "
        "must clearly state it is triage support, not a diagnosis. Several interviewees also expressed interest in capturing cough and sputum "
        "in one session and tracking referrals to GeneXpert or clinical workup. These insights directly shaped the researchers' decision to "
        "include staff review, multimodal fusion, referral tracking, and prominent disclaimers in TBhon."
    )

    add_heading(doc, "Survey/Questionnaire", 4)
    add_para(
        doc,
        "A structured questionnaire was distributed to a broader group of respondents to quantify challenges identified during interviews. "
        "The survey covered topics such as familiarity with TB symptoms, willingness to use mobile-assisted screening, importance of sputum "
        "and cough capture, and perceived usefulness of IoT hardware in booth settings. The questionnaire was administered digitally to reach "
        "respondents within health-related academic and community networks in Davao City."
    )
    add_para(
        doc,
        "The researchers observed that a significant majority of respondents considered integrated digital screening valuable for booth and "
        "RHU settings, particularly when staff remain in control of results. Respondents also indicated preference for mobile-based workflows "
        "with session history and PDF export, confirming the need for the integrated functionalities proposed in TBhon."
    )

    add_heading(doc, "Observation", 4)
    add_para(
        doc,
        "Observational methods supplemented interview and survey data by examining how staff interact with existing screening workflows. "
        "The researchers observed health education sessions and mock booth setups, noting workflow inefficiencies such as switching between "
        "paper forms, separate recording apps, and manual referral notes."
    )
    add_para(
        doc,
        "Through observation, the researchers noticed that staff prefer minimal interaction steps during client-facing screening and clear "
        "visual feedback during ML processing. Observations also revealed that patients benefit from a simple QR-based way to access results "
        "after leaving the booth, which guided the group's decisions on patient claim flows and result screen design."
    )

    add_heading(doc, "Requirements Based on User Perspectives", 4)

    add_para(doc, "User Requirements", bold=True)
    add_bullets(
        doc,
        [
            "Staff must be able to register with a facility invite code and securely log in before conducting screenings.",
            "Staff must be able to capture client demographics and complete an 11-question TB symptom and risk checklist.",
            "Staff must be able to record three cough audio samples via phone microphone or IoT hardware, with quality feedback before upload.",
            "Staff must be able to capture or skip sputum smear images with documented skip reasons.",
            "Staff must receive a fused low/moderate/high triage risk score with modality breakdown and must confirm results before client disclosure.",
            "Staff must be able to view session history, update referral status, and export a PDF triage summary.",
            "Patients must be able to claim screening results via QR code or deep link after staff completes the session.",
            "Administrators must be able to manage facilities and assign staff to screening booths.",
        ],
    )

    add_para(doc, "Functional Requirements", bold=True)
    add_bullets(
        doc,
        [
            "The system shall integrate Expo/React Native for the mobile staff and patient interfaces.",
            "The system shall use an Express/Prisma backend with MySQL for authentication, session storage, and media persistence.",
            "The system shall implement a FastAPI inference service with PyTorch CNN models for cough Mel-spectrogram and sputum image classification.",
            "The system shall perform cough quality checks (/check-quality) and sputum quality checks before inference.",
            "The system shall fuse checklist, cough ML, and sputum ML outputs via weighted log-odds risk stratification.",
            "The system shall support ESP32 IoT capture through BLE Wi‑Fi setup and backend IoT polling/upload endpoints.",
            "The system shall generate patient access tokens and QR codes for result claiming.",
            "The system shall provide in-app TB education content and PDF export with triage-only disclaimers.",
        ],
    )

    add_para(doc, "Data Requirements", bold=True)
    add_bullets(
        doc,
        [
            "User account data: names, email, hashed credentials, role (STAFF/ADMIN/PATIENT), facility linkage, and verification status.",
            "Screening session data: client demographics, checklist answers, staff operator, facility, timestamps, and referral status.",
            "Cough recording data: raw audio bytes, quality check results, TB audio prediction probabilities, and recording metadata.",
            "Sputum image data: image bytes, quality check results, phlegm/AFB prediction outputs, and skip reasons when applicable.",
            "Screening result data: fused risk level (low/moderate/high), recommendation text, staff notes, and patient access tokens.",
            "IoT device data: device identifiers, presence status, queued commands, and uploaded media references.",
        ],
    )

    add_para(doc, "Environmental Requirements", bold=True)
    add_bullets(
        doc,
        [
            "The application must function on Android and iOS smartphones via Expo, with native builds required for BLE/ESP32 features.",
            "The system requires a stable internet connection for ML inference, backend synchronization, and Cloudflare tunnel access.",
            "The mobile device must support microphone and camera hardware for phone-only capture workflows.",
            "The backend and ML services must operate reliably on cloud infrastructure (DigitalOcean droplets with PM2/systemd).",
            "IoT modules must operate in booth environments with Wi‑Fi connectivity to upload captures to the backend.",
        ],
    )

    add_para(doc, "Usability Requirements", bold=True)
    add_bullets(
        doc,
        [
            "The interface must be learnable by booth staff without prior ML or medical informatics training.",
            "Screening steps must follow a clear linear flow with back-navigation and visible progress indicators.",
            "Triage disclaimers must be prominent on result and PDF screens.",
            "Error messages and loading states must clearly indicate ML processing, upload status, and connectivity failures.",
            "The interface must remain readable on common mobile screen sizes used in field settings.",
        ],
    )

    add_para(doc, "Designer Requirements", bold=True)
    add_bullets(
        doc,
        [
            "The system architecture must be modular (mobile, backend, ML, IoT firmware) to allow independent development and deployment.",
            "The codebase must follow clean code principles with OpenAPI/Swagger documentation for backend APIs.",
            "ML model paths, API URLs, and secrets must be configurable through environment variables.",
            "The design must support iterative refinement based on model calibration, heuristic evaluation, and UAT feedback.",
        ],
    )

    add_heading(doc, "C. Storyboarding and Prototyping", 3)
    add_para(
        doc,
        "Storyboarding and Prototyping is the stage in the UCSD process where design ideas are translated into tangible representations "
        "that stakeholders and users can visualize and evaluate. The research group developed an interaction storyboard illustrating the "
        "end-to-end staff screening journey, followed by screen-level prototypes for the major interfaces of TBhon."
    )

    add_heading(doc, "Storyboard", 4)
    add_para(
        doc,
        "The storyboard presents the narrative of how a typical booth staff member conducts a walk-in TB screening. The story follows a "
        "staff operator who greets a client, records intake information, administers the symptom checklist, guides three cough recordings, "
        "captures a sputum smear photo (or documents a skip reason), waits for ML processing, reviews the fused triage risk on a staff-review "
        "screen, explains the result with counseling support from the Learn module, displays a QR code for the client to claim the result, "
        "and marks referral to GeneXpert when risk is moderate or high."
    )
    add_para(
        doc,
        "During storyboarding, the group identified an important gap in an early draft: there was no dedicated staff-review step before "
        "showing the client the triage score. This was added to reinforce human oversight and align with the system's honest-scope positioning."
    )

    add_heading(doc, "Prototype", 4)
    add_para(doc, "The prototypes cover the following key screens and interaction flows:")
    screens = [
        ("Home / Dashboard Screen", "Displays quick actions to start screening, session history, and role-appropriate navigation for staff or patient."),
        ("Client Intake Screen", "Captures screened person demographics, contact details, and optional identification information."),
        ("Checklist Screen", "Presents 11 canonical symptom and risk questions with explanatory subtext for staff-assisted completion."),
        ("Cough Recording Screen", "Guides three cough captures with timer, quality feedback, and upload confirmation (phone or IoT path)."),
        ("Sputum Capture Screen", "Supports camera capture, quality check, skip-reason modal, and preview of smear image."),
        ("Processing Screen", "Shows ML inference progress for cough, sputum, and risk fusion with modality breakdown."),
        ("Staff Review Screen", "Requires staff confirmation and optional notes before revealing triage risk to the client."),
        ("Result Screen", "Displays low/moderate/high risk, referral guidance, patient QR, PDF export, and triage disclaimers."),
        ("History and Details Screens", "List past sessions; replay audio, view checklist answers, sputum images, and referral status."),
        ("Patient Claim Screens", "QR/deep-link preview and account linking for screened persons to access their results."),
        ("Admin Facilities Screen", "Facility CRUD and staff assignment for multi-booth deployment."),
    ]
    for title, desc in screens:
        p = doc.add_paragraph()
        r1 = p.add_run(title + " — ")
        r1.bold = True
        r2 = p.add_run(desc)
        r1.font.size = r2.font.size = Pt(12)

    add_para(
        doc,
        "The researchers developed initial prototypes using Figma, enabling linked wireframes that simulated realistic staff screening flows. "
        "Each screen was annotated with expected interactions and system responses. During prototype development, the group iterated on information "
        "density on the result screen—balancing modality breakdown detail with clear triage messaging and disclaimer visibility."
    )

    add_heading(doc, "D. Evaluation of Prototype", 3)
    add_para(
        doc,
        "Prototype evaluation assessed the proposed design systematically to identify usability problems and areas for improvement before "
        "full implementation. The research group employed Heuristic Evaluation as the primary method, selecting the best design among three "
        "alternative interface concepts and evaluating it against Nielsen's Ten Usability Heuristics (Nielsen, 1994)."
    )

    add_heading(doc, "Design Selection Process", 4)
    add_para(
        doc,
        "Each member developed an alternative design concept varying navigation layout, color scheme, and information hierarchy. "
        "Design Alternative A featured a bottom navigation bar with card-based screening steps. "
        "Design Alternative B used a top drawer with a split checklist/recording layout. "
        "Design Alternative C emphasized a minimalist stepper with floating action buttons. "
        "After deliberation, Design Alternative A was selected for heuristic evaluation due to its intuitive flow, consistency with familiar "
        "mobile health app patterns, and clear visual hierarchy for booth staff use."
    )

    add_heading(doc, "Heuristic Evaluation Format", 4)
    add_para(doc, "Table 2.2. Heuristic Evaluation of TBhon Prototype", bold=True)
    add_table(
        doc,
        ["#", "Heuristic", "Observed Strengths", "Issues Identified", "Severity (1–4)"],
        [
            [
                "1",
                "Visibility of System Status",
                "Processing screen shows ML steps; IoT timeline displays device presence and upload status.",
                "Some ML inference transitions lack estimated time remaining during slow network conditions.",
                "2 – Minor",
            ],
            [
                "2",
                "Match Between System and Real World",
                "Checklist language mirrors DOH-style symptom questions; referral line mentions GeneXpert/clinical workup.",
                "Some modality breakdown labels (e.g., log-odds fusion) use technical terms unfamiliar to booth staff.",
                "2 – Minor",
            ],
            [
                "3",
                "User Control and Freedom",
                "Staff can navigate back, skip sputum with documented reason, and cancel IoT capture.",
                "No confirmation before completing a session when checklist answers are incomplete.",
                "3 – Major",
            ],
            [
                "4",
                "Consistency and Standards",
                "Icons, triage colors, and disclaimer copy are consistent across result and PDF export.",
                "Minor font size differences between history cards and checklist detail view.",
                "1 – Cosmetic",
            ],
            [
                "5",
                "Error Prevention",
                "Cough quality gate blocks poor recordings; sputum quality check prompts retake.",
                "No guardrail when staff attempts to finalize without staff-review confirmation in early prototype.",
                "3 – Major",
            ],
            [
                "6",
                "Recognition Rather Than Recall",
                "Screening step order is visible; session history shows client name, date, and risk level.",
                "IoT Wi‑Fi setup steps require recall of device LED states without persistent on-screen legend.",
                "2 – Minor",
            ],
            [
                "7",
                "Flexibility and Efficiency of Use",
                "Walk-in shortcut creates draft session; phone-only path bypasses IoT setup.",
                "No quick-repeat screening template for returning clients in the prototype.",
                "2 – Minor",
            ],
            [
                "8",
                "Aesthetic and Minimalist Design",
                "Result screen prioritizes triage level and disclaimer; Learn content is separated from screening flow.",
                "Staff review screen can feel dense when all three modality scores are expanded.",
                "2 – Minor",
            ],
            [
                "9",
                "Help Users Recognize, Diagnose, and Recover From Errors",
                "Upload failures show retry options; auth errors distinguish network vs credential issues.",
                "Generic message when ML service is unreachable does not guide staff to phone-only fallback steps.",
                "3 – Major",
            ],
            [
                "10",
                "Help and Documentation",
                "In-app Learn module explains TB symptoms and when to seek care; intake includes booth copy.",
                "No dedicated FAQ for IoT troubleshooting during first-time booth setup.",
                "2 – Minor",
            ],
        ],
    )
    add_para(
        doc,
        "Severity ratings were assigned using Nielsen's four-point scale: 1 – Cosmetic, 2 – Minor, 3 – Major, 4 – Usability catastrophe. "
        "The most significant issues addressed before full development were incomplete-checklist confirmation, mandatory staff-review guardrails, "
        "and improved ML connectivity error guidance. Overall, the evaluation confirmed strong alignment with usability best practices in "
        "consistency, triage disclaimer visibility, and staff-guided workflow design."
    )

    doc.add_page_break()

    # --- CHAPTER III ---
    add_heading(doc, "CHAPTER III", 1)
    add_heading(doc, "CONCLUSION AND RECOMMENDATION", 2)

    add_heading(doc, "Conclusion", 3)
    conclusions = [
        (
            "The development of TBhon: An IoT-Enabled Mobile-Based Multimodal Diagnostic Support System for Tuberculosis Pre-screening "
            "demonstrated how Human-Computer Interaction (HCI) principles can be applied to create a system that is user-centered, functional, "
            "and responsive to real-world TB screening challenges in resource-limited settings. Through the User-Centered System Design (UCSD) "
            "process, the researchers identified the needs of booth staff, patients, and administrators and transformed them into practical "
            "features and interface designs."
        ),
        (
            "The study highlighted that existing TB screening approaches often rely on fragmented tools, single-modality analysis, or "
            "infrastructure-heavy laboratory workflows while giving limited attention to integrated multimodal triage, staff oversight, and "
            "honest-scope mobile delivery. TBhon addressed these limitations by combining an 11-question checklist, CNN-based cough audio "
            "analysis (Mel-spectrogram + hybrid CNN/GBM), sputum smear image classification, weighted log-odds risk fusion, optional ESP32 "
            "IoT capture, cloud-backed session management, patient QR claiming, and referral tracking within a single staff-operated platform."
        ),
        (
            "The evaluation results showed that the selected prototype aligned well with established usability principles, particularly in "
            "system status visibility during ML processing, consistent triage disclaimer language, and a linear staff-guided screening flow. "
            "Although heuristic evaluation identified concerns such as incomplete-session guardrails and IoT setup recall, these findings "
            "guided prototype refinement before full implementation."
        ),
        (
            "From an HCI perspective, the project emphasized designing for booth workflows rather than algorithmic performance alone. Effective "
            "interface design required balancing clinical honesty (triage support, not diagnosis), staff control (review before disclosure), "
            "patient access (QR claim), and technical modularity (mobile, backend, ML, IoT). Interviews, surveys, observations, task analysis, "
            "storyboarding, and heuristic evaluation helped the group understand how design decisions influence trust, efficiency, and "
            "follow-up behavior in community screening contexts."
        ),
        (
            "Overall, the study demonstrated that TBhon has the potential to support earlier TB risk identification and more organized booth "
            "screening in Davao City and similar settings. The project also strengthened the researchers' understanding of user-centered design, "
            "multimodal health AI integration, and real-world mobile system development aligned with the Philippine Strategic TB Elimination Plan."
        ),
    ]
    for c in conclusions:
        add_para(doc, c)

    add_heading(doc, "Recommendation", 3)
    add_para(
        doc,
        "Based on the findings and limitations of the study, the researchers recommend the following improvements and future enhancements "
        "for the continued development of TBhon."
    )
    recommendations = [
        (
            "First,",
            "future developers may conduct formal User Acceptance Testing (UAT) with a larger pool of community health workers, nurses, and "
            "IT experts using ISO/IEC 25010–aligned instruments to validate functional suitability, performance efficiency, usability, reliability, "
            "and security in live booth settings.",
        ),
        (
            "Second,",
            "the ML pipeline may be expanded with prospective field validation using institution-approved clinical data, external test sets, "
            "and calibration on local recording conditions to improve generalization beyond publicly available Kaggle and repository datasets.",
        ),
        (
            "Third,",
            "the application may support limited offline or degraded-mode operation—caching checklist and intake data locally when connectivity "
            "is unstable and synchronizing when the tunnel/backend connection is restored, improving reliability in rural screening sites.",
        ),
        (
            "Fourth,",
            "the IoT module may be enhanced with clearer on-device status indicators, in-app troubleshooting guides, and automated firmware "
            "update workflows to reduce setup friction for first-time booth deployment.",
        ),
        (
            "Fifth,",
            "future researchers may integrate additional modalities (e.g., chest X-ray CAD interfaces or GeneXpert result logging) only as "
            "supplementary triage signals—maintaining honest scope and avoiding presentation as standalone diagnosis.",
        ),
        (
            "Sixth,",
            "multilingual support (e.g., Cebuano/Bisaya and Filipino) for checklist, Learn content, and result counseling may improve "
            "accessibility for diverse clients in Mindanao screening booths.",
        ),
        (
            "Finally,",
            "future HCI studies may use TBhon as a reference for examining trust, disclaimer comprehension, and staff–patient interaction in "
            "AI-assisted public health triage—contributing to smart health initiatives and community-oriented TB elimination programs.",
        ),
    ]
    for lead, body in recommendations:
        p = doc.add_paragraph()
        r1 = p.add_run(lead + " ")
        r1.bold = True
        r2 = p.add_run(body)
        r1.font.size = r2.font.size = Pt(12)

    doc.save(OUT)
    print("Saved:", OUT)


if __name__ == "__main__":
    build()
