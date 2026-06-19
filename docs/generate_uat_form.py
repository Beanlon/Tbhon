"""Generate TBhon User Acceptance Testing (UAT) form — ISO/IEC 25010 aligned."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = r"C:\Users\Mika\Documents\GitHub\Tbhon\docs\TBhon_UAT_Form.docx"

# (criterion, ISO/IEC 25010 characteristic)
CRITERIA = [
    (
        "The system supports the complete TB pre-screening workflow "
        "(client intake, checklist, cough capture, sputum capture, staff review, and result disclosure).",
        "Functional Suitability",
    ),
    (
        "The 11-item symptom checklist captures relevant TB screening indicators for triage.",
        "Functional Suitability",
    ),
    (
        "Cough audio analysis (CNN / Mel-spectrogram) produces useful and interpretable triage signals.",
        "Functional Suitability",
    ),
    (
        "Sputum smear image analysis produces useful and interpretable triage signals.",
        "Functional Suitability",
    ),
    (
        "Multimodal risk fusion (Low, Moderate, High) is appropriate for preliminary TB triage support.",
        "Functional Suitability",
    ),
    (
        "Screening reports, recommendations, and referral guidance are clear and actionable.",
        "Functional Suitability",
    ),
    (
        "Patient result access (QR claim) and screening history functions work as intended.",
        "Functional Suitability",
    ),
    (
        "Cough audio classification completes within an acceptable time (target: ≤ 5 seconds).",
        "Performance Efficiency",
    ),
    (
        "Sputum image classification completes within an acceptable time (target: ≤ 5 seconds).",
        "Performance Efficiency",
    ),
    (
        "The mobile application responds smoothly during routine screening tasks.",
        "Performance Efficiency",
    ),
    (
        "ML inference and cloud backend connectivity perform adequately during testing.",
        "Performance Efficiency",
    ),
    (
        "The application is easy to use for booth staff with minimal technical training.",
        "Usability",
    ),
    (
        "The interface design is clear, consistent, and appropriate for a health screening context.",
        "Usability",
    ),
    (
        "Navigation through the staff-guided screening flow is intuitive and logical.",
        "Usability",
    ),
    (
        "Triage disclaimers (not a diagnosis) are visible, understandable, and appropriately placed.",
        "Usability",
    ),
    (
        "Processing status and error guidance during ML analysis are clear and helpful.",
        "Usability",
    ),
    (
        "The system operates reliably without unexpected crashes or data loss during screening.",
        "Reliability",
    ),
    (
        "Screening session data is stored correctly and can be retrieved from history.",
        "Reliability",
    ),
    (
        "ESP32 IoT cough/sputum capture integrates reliably when used (if tested).",
        "Reliability",
    ),
    (
        "The system handles unstable network conditions in a predictable manner.",
        "Reliability",
    ),
    (
        "User authentication (login, registration, password policy) protects system access appropriately.",
        "Security",
    ),
    (
        "Patient screening data is handled with appropriate privacy and confidentiality.",
        "Security",
    ),
    (
        "Cloud-backed data storage (Express/Prisma backend) is trustworthy for screening records.",
        "Security",
    ),
    (
        "Overall, I am satisfied with TBhon as a TB pre-screening and triage support tool.",
        "Overall Satisfaction",
    ),
]

UAT_TASKS = [
    "Register or log in as booth staff and start a new walk-in screening session.",
    "Complete client intake and the 11-item TB symptom checklist.",
    "Record or upload cough audio (mobile mic or ESP32 IoT module, if available).",
    "Capture or upload a sputum smear image (camera or ESP32 microscope module, if available).",
    "Review ML processing status, staff-review screen, and fused triage risk result (Low / Moderate / High).",
    "Verify screening history, patient QR result claim (if applicable), and disclaimer visibility.",
]


def shade_cell(cell, fill_hex: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_section_heading(doc, letter: str, title: str):
    p = doc.add_paragraph()
    run = p.add_run(f"Section {letter}: {title}")
    run.bold = True
    run.font.size = Pt(12)


def add_info_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], label, bold=True)
        set_cell_text(table.rows[i].cells[1], value)
        shade_cell(table.rows[i].cells[0], "E8F5E9")
    doc.add_paragraph()


def add_uat_header(doc):
    table = doc.add_table(rows=2, cols=1)
    table.autofit = True
    banner = table.rows[0].cells[0]
    shade_cell(banner, "1B5E20")
    set_cell_text(
        banner,
        "USER ACCEPTANCE TESTING (UAT) FORM",
        bold=True,
        size=16,
        color=RGBColor(255, 255, 255),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    subtitle = table.rows[1].cells[0]
    shade_cell(subtitle, "2E7D32")
    set_cell_text(
        subtitle,
        "TBhon: IoT-Enabled Mobile-Based Multimodal TB Pre-Screening System",
        bold=True,
        size=12,
        color=RGBColor(255, 255, 255),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()


def add_criteria_table(doc):
    headers = ["No.", "Criteria", "ISO/IEC 25010", "5", "4", "3", "2", "1", "Remarks"]
    table = doc.add_table(rows=1 + len(CRITERIA), cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell, "C8E6C9")

    for idx, (criterion, iso_char) in enumerate(CRITERIA, start=1):
        row = table.rows[idx]
        set_cell_text(row.cells[0], str(idx), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], criterion, size=9)
        set_cell_text(row.cells[2], iso_char, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        for col in range(3, 8):
            set_cell_text(row.cells[col], "", size=10)
        set_cell_text(row.cells[8], "", size=9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Rating Scale: ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "5 – Strongly Agree   |   4 – Agree   |   3 – Neutral   |   "
        "2 – Disagree   |   1 – Strongly Disagree"
    )
    run.font.size = Pt(10)


def add_overall_rating(doc):
    add_section_heading(doc, "D", "Overall Rating")
    p = doc.add_paragraph("How would you rate the overall performance of TBhon?")
    for run in p.runs:
        run.font.size = Pt(11)

    options = ["☐ Excellent", "☐ Good", "☐ Fair", "☐ Poor", "☐ Very Poor"]
    table = doc.add_table(rows=1, cols=len(options))
    table.style = "Table Grid"
    for i, opt in enumerate(options):
        set_cell_text(table.rows[0].cells[i], opt, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()


def add_comments_section(doc):
    add_section_heading(doc, "E", "Comments / Suggestions")
    doc.add_paragraph(
        "Please provide any additional feedback on usability, accuracy, reliability, security, "
        "or features that could improve TBhon for community TB screening booths."
    )
    for _ in range(5):
        p = doc.add_paragraph("_" * 90)
        for run in p.runs:
            run.font.size = Pt(11)
    doc.add_paragraph()


def add_signature_section(doc):
    add_section_heading(doc, "F", "Evaluator Confirmation")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "Signature:", bold=True)
    set_cell_text(table.rows[0].cells[1], "Date:")
    set_cell_text(table.rows[1].cells[0], "_" * 40)
    set_cell_text(table.rows[1].cells[1], "_" * 25)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_uat_header(doc)

    p = doc.add_paragraph()
    run = p.add_run(
        "Instructions: "
    )
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "Complete this form after performing the UAT tasks below using the TBhon mobile application. "
        "Rate each criterion based on your professional experience. One form per evaluator."
    )
    run.font.size = Pt(10)
    doc.add_paragraph()

    add_section_heading(doc, "A", "Evaluator Information")
    add_info_table(
        doc,
        [
            ("Name:", "_______________________________________________"),
            ("Role / Position:", "_______________________________________________"),
            (
                "Organization:",
                "_______________________________________________",
            ),
            ("Date:", "_______________________________________________"),
        ],
    )

    add_section_heading(doc, "B", "System Information")
    add_info_table(
        doc,
        [
            ("System Name:", "TBhon"),
            (
                "Version:",
                "1.0 (Prototype / Field Evaluation Build)",
            ),
            (
                "Purpose:",
                "To support tuberculosis pre-screening and triage using an 11-item symptom checklist, "
                "CNN-based cough audio analysis (Mel-spectrogram), sputum smear image classification, "
                "weighted multimodal risk fusion (Low / Moderate / High), optional ESP32 IoT capture, "
                "and cloud-backed session management via React Native mobile app, Express/Prisma backend, "
                "and FastAPI ML inference service.",
            ),
            (
                "Scope Note:",
                "TBhon is triage support only — not a medical diagnosis. Standard clinical evaluation "
                "and laboratory testing are required for diagnosis and treatment decisions.",
            ),
        ],
    )

    add_section_heading(doc, "B.1", "UAT Tasks (complete before rating)")
    for i, task in enumerate(UAT_TASKS, start=1):
        p = doc.add_paragraph(f"{i}. {task}", style="List Number")
        for run in p.runs:
            run.font.size = Pt(10)
    doc.add_paragraph()

    add_section_heading(doc, "C", "Evaluation Criteria (ISO/IEC 25010 Software Quality Model)")
    add_criteria_table(doc)
    doc.add_paragraph()

    add_overall_rating(doc)
    add_comments_section(doc)
    add_signature_section(doc)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Target Respondents: ")
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(
        "Community health workers, nurses, public health professionals, and IT experts "
        "(purposive sampling per study methodology)."
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
