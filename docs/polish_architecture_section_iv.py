"""Rebuild Section IV with manuscript figures and styled tables."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "docs" / "figures"
SRC = ROOT / "docs" / "JAM_Architecture_Paper.docx"
OUT = ROOT / "docs" / "JAM_Architecture_Paper_updated.docx"
BACKUP = ROOT / "docs" / "JAM_Architecture_Paper.backup.docx"
DL_OUT = Path(r"C:\Users\Mika\Downloads\JAM_Architecture Paper_updated.docx")

HEADER_FILL = "D6EAF8"  # light blue header band (matches chapter/demo tables)


def shade_cell(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_table(table: Table) -> None:
    table._element.getparent().remove(table._element)


def insert_table_after(paragraph: Paragraph, rows: int, cols: int) -> Table:
    doc = paragraph.part.document
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    paragraph._p.addnext(table._tbl)
    return table


def fill_table(table: Table, headers: list[str], data: list[list[str]]) -> None:
    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        shade_cell(cell, HEADER_FILL)
    for row_idx, row in enumerate(data, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def insert_figure_after(paragraph: Paragraph, image_path: Path, caption: str, width_in: float = 5.5) -> Paragraph:
    cursor = paragraph
    if image_path.is_file():
        img_p = insert_paragraph_after(cursor)
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.add_run().add_picture(str(image_path), width=Inches(width_in))
        cursor = img_p
    else:
        cursor = insert_paragraph_after(cursor, f"[Insert figure: {image_path.name}]")
    cap_p = insert_paragraph_after(cursor, caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap_p.runs:
        cap_p.runs[0].bold = True
        cap_p.runs[0].font.size = Pt(10)
    return cap_p


# (style, text) | ("table", key) | ("figure", path, caption, width)
SECTION_BLOCKS: list[tuple] = [
    (
        "Body Text",
        "This section reports datasets, experimental apparatus, training procedures, and quantitative findings "
        "for the TBhon multimodal tuberculosis pre-screening system. Evaluation encompassed unimodal cough and "
        "sputum classifiers, three-fold cross-validation, architectural ablation (convolutional baseline versus "
        "hybrid CNN+gradient-boosted ensemble), operating-point calibration, pre-inference cough quality control, "
        "and end-to-end API response behavior. Screening-level multimodal risk integration is described "
        "qualitatively; a dedicated held-out evaluation of the full fused triage score against microbiological "
        "ground truth was not completed and is noted as a limitation in Section IV-G.",
    ),
    ("Heading 2", "A. Datasets and Experimental Setup"),
    ("Body Text", "The evaluation employed two publicly available data sources aligned with the TBhon training pipeline."),
    (
        "Body Text",
        "Cough audio. The open Kaggle tuberculosis audio corpus provided 7,817 labeled respiratory recordings. "
        "All waveforms were converted to mono, resampled to 16 kHz, and normalized to a fixed temporal window "
        "before feature extraction. The deployed hybrid classifier used four-second clips and sixty-four Mel "
        "frequency bins. Labels were binary (tuberculosis-positive versus tuberculosis-negative). The publisher "
        "supplied three predefined cross-validation folds, each allocating approximately 5,211 recordings to "
        "training and 2,606 to testing.",
    ),
    (
        "Body Text",
        "Sputum smear images. A curated set of 1,438 Ziehl-Neelsen stained microscopy fields was partitioned "
        "into 1,149 training, 142 validation, and 147 test images (approximately 80% / 10% / 10%). "
        "Acid-fast bacillus (AFB) labels were derived from bounding-box annotations by counting visible bacilli "
        "per field. Binary AFB classification (production path) used 224 x 224-pixel ResNet-18 inputs.",
    ),
    (
        "Body Text",
        "Prototype acquisition used an ESP32-S3 with INMP441 microphone and microscope-mounted camera (640 x 480 JPEG). "
        "Compute environment: Intel Core i5, 8 GB RAM, CUDA GPU, PyTorch 2.x; production inference on a DigitalOcean "
        "FastAPI ML droplet (/check-quality, /predict, /predict-phlegm).",
    ),
    ("Heading 2", "B. Cough Audio Classification"),
    (
        "Body Text",
        "Two cough classifiers were compared on held-out fold 0 (n = 2,606). The baseline was a CNN on log-Mel "
        "spectrograms (128 bins, six-second clips), trained with AdamW (3 x 10^-4), batch size 32, and early stopping. "
        "The production hybrid combined a CNN branch (64 Mel bins, four-second clips, eight epochs) with a "
        "gradient-boosted ensemble on MFCC and summary features, fused by validation-tuned weighted averaging.",
    ),
    (
        "Body Text",
        "Table I summarizes the fold-0 ablation. The hybrid model improved macro-F1 by 17.8 percentage points over "
        "the convolutional baseline (50.6% to 68.3%).",
    ),
    ("table head", "TABLE I\nCOUGH CLASSIFIER ABLATION ON FOLD 0 (HELD-OUT TEST, n = 2,606)"),
    (
        "figure",
        FIG / "figure_5_4_cough_architecture_and_confusion_matrix.png",
        "Figure VII. Hybrid CNN + GBM Cough Classification Architecture and Confusion Matrix",
        5.8,
    ),
    (
        "figure",
        FIG / "cough_metrics_compilation_8_epochs.png",
        "Figure VIII. Hybrid Cough CNN Metrics Compilation at 8 Epochs",
        5.8,
    ),
    ("Heading 2", "C. Cross-Fold Generalization"),
    (
        "Body Text",
        "The hybrid architecture was independently trained and evaluated on all three cross-validation folds (Table II). "
        "Mean macro-F1 was 68.6% (range 68.3-69.1 pp). Fold 1 achieved the highest macro-F1 (69.1%) and was "
        "selected for production deployment.",
    ),
    ("table head", "TABLE II\nHYBRID COUGH CLASSIFIER CROSS-FOLD TEST PERFORMANCE"),
    ("Heading 2", "D. Sputum Smear Classification"),
    (
        "Body Text",
        "The binary AFB classifier used a ResNet-18 backbone fine-tuned for thirty epochs (AdamW, 1 x 10^-3, "
        "class-weighted loss). The decision threshold was tuned on validation data under a >=95% AFB-positive "
        "sensitivity constraint (production threshold = 0.63).",
    ),
    (
        "Body Text",
        "On the held-out test partition (n = 216; 6 AFB-negative, 210 AFB-positive), the production model "
        "(phlegm_afb_binary_20260531_133949) achieved 94.91% accuracy, 66.32% macro-F1, and 96.19% AFB-positive "
        "sensitivity (Table III). AFB-negative specificity was 50.00% (3 TN, 3 FP), reflecting screening-oriented "
        "thresholding and the small negative test cohort.",
    ),
    ("table head", "TABLE III\nPRODUCTION UNIMODAL CLASSIFIER PERFORMANCE ON HELD-OUT TEST DATA"),
    (
        "figure",
        FIG / "sputum_metrics_compilation_30_epochs.png",
        "Figure IX. Sputum ResNet18 Binary Classifier Metrics Compilation at 30 Epochs",
        5.8,
    ),
    (
        "figure",
        FIG / "sputum_confusion_matrix.png",
        "Figure X. Sputum AFB Binary Classifier Confusion Matrix",
        3.2,
    ),
    ("Heading 2", "E. Pre-Inference Quality Control"),
    (
        "Body Text",
        "Each cough recording passes a heuristic /check-quality gate (RMS, crest factor, spectral flatness, "
        "periodicity, dynamic range) before classification. Invalid clips are excluded from fusion. Functional "
        "testing confirmed rejection of noise, silence, and speech-like inputs while accepting cough bursts.",
    ),
    ("Heading 2", "F. Multimodal Screening Fusion"),
    (
        "Body Text",
        "End-to-end risk scoring integrates checklist probability (weight 0.85), mean cough ML probability (1.00), "
        "and sputum ML probability (0.70) through weighted log-odds fusion: "
        "P_fused = sigmoid( SUM_i w_i * logit(p_i) / SUM_i w_i ). Clinical safety floors elevate high-concern "
        "checklist or confident AFB-positive findings. Outputs map to Low (<0.38), Moderate (0.38-0.62), or High "
        "(>=0.62) risk (Table IV).",
    ),
    ("table head", "TABLE IV\nMULTIMODAL FUSION MODALITY RELIABILITY WEIGHTS"),
    (
        "figure",
        FIG / "fusion_diagram_checklist_cough_sputum.png",
        "Figure XI. Multimodal Fusion Diagram (Checklist + Cough + Sputum)",
        5.8,
    ),
    (
        "figure",
        FIG / "multimodal_fusion_result_mockup.png",
        "Figure XII. Sample Multimodal Fusion and Triage Result Output",
        4.2,
    ),
    ("Heading 2", "G. Evaluation Metrics, System Response, and Discussion"),
    (
        "Body Text",
        "Metrics: accuracy, precision, recall, and macro-F1 on held-out test partitions. Response time was evaluated "
        "against the five-second per-inference target; four of six UAT evaluators rated inference speed as Agree or "
        "Strongly Agree under stable Wi-Fi, with occasional delays on unstable networks.",
    ),
    (
        "Body Text",
        "Production cough hybrid: 75.8% accuracy, 69.1% macro-F1, 88.3% non-TB recall, 47.6% TB recall. Sputum: "
        "94.9% accuracy, 96.2% AFB+ sensitivity. All six end-to-end workflow test cases passed on iPhone 15 and "
        "Redmi Note 12 Pro 5G. Limitations: public datasets only, low cough TB recall for standalone use, unstable "
        "sputum specificity estimates (n=6 negatives), no held-out fused-score benchmark, observational latency only.",
    ),
]

TABLE_SPECS: dict[str, tuple[list[str], list[list[str]]]] = {
    "TABLE I": (
        ["Model", "Fold", "Test n", "Accuracy", "Precision", "Recall", "Macro-F1"],
        [
            ["Convolutional network (baseline)", "0", "2,606", "50.6%", "60.4%", "60.7%", "50.6%"],
            ["Hybrid CNN+GBM (ablation)", "0", "2,606", "73.9%", "68.3%", "68.4%", "68.3%"],
        ],
    ),
    "TABLE II": (
        ["Fold", "Test n", "Accuracy", "Macro-F1", "TB Recall", "Non-TB Recall"],
        [
            ["0", "2,606", "73.9%", "68.3%", "55.2%", "81.5%"],
            ["1 (production)", "2,606", "75.8%", "69.1%", "47.6%", "88.3%"],
            ["2", "2,605", "76.0%", "68.3%", "44.0%", "90.0%"],
            ["Mean", "N/A", "75.2%", "68.6%", "48.9%", "86.6%"],
        ],
    ),
    "TABLE III": (
        ["Production Model", "Test n", "Accuracy", "Macro Precision", "Macro Recall", "Macro-F1"],
        [
            ["Hybrid cough CNN+GBM (fold 1)", "2,606", "75.8%", "71.7%", "67.9%", "69.1%"],
            ["ResNet-18 AFB binary (sputum)", "216", "94.9%", "62.6%", "73.0%", "66.3%"],
            ["AFB+ sensitivity (sputum)", "210 positives", "N/A", "98.5%", "96.2%", "97.1%"],
        ],
    ),
    "TABLE IV": (
        ["Modality", "Reliability Weight"],
        [
            ["Symptom and exposure checklist", "0.85"],
            ["Cough machine learning probability", "1.00"],
            ["Sputum machine learning probability", "0.70"],
        ],
    ),
}


def main() -> None:
    if not BACKUP.is_file():
        shutil.copy2(SRC, BACKUP)

    doc = Document(str(BACKUP))

    while doc.tables:
        remove_table(doc.tables[0])

    start_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "IV. EXPERIMENTAL RESULTS")
    end_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "CONCLUSION AND RECOMMENDATION")

    anchor = doc.paragraphs[start_idx]
    for idx in range(end_idx - 1, start_idx, -1):
        delete_paragraph(doc.paragraphs[idx])

    cursor = anchor
    for block in SECTION_BLOCKS:
        if block[0] == "figure":
            _, path, caption, width = block
            cursor = insert_figure_after(cursor, path, caption, width)
            spacer = OxmlElement("w:p")
            cursor._p.addnext(spacer)
            cursor = Paragraph(spacer, cursor._parent)
            continue

        style, text = block
        cursor = insert_paragraph_after(cursor, text, style)
        if style == "table head":
            key = text.split("\n", 1)[0].strip()
            headers, data = TABLE_SPECS[key]
            table = insert_table_after(cursor, 1 + len(data), len(headers))
            fill_table(table, headers, data)
            spacer = OxmlElement("w:p")
            table._tbl.addnext(spacer)
            cursor = Paragraph(spacer, cursor._parent)

    doc.save(str(OUT))
    shutil.copy2(OUT, DL_OUT)
    print("Polished:", OUT)
    print("Downloads copy:", DL_OUT)


if __name__ == "__main__":
    main()
