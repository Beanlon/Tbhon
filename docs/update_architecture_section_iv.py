"""Replace Section IV (EXPERIMENTAL RESULTS) in JAM_Architecture Paper.docx."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table

SRC = Path(r"C:\Users\Mika\Downloads\JAM_Architecture Paper.docx")
BACKUP = SRC.with_name("JAM_Architecture Paper.backup.docx")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_cell(table, row: int, col: int, value: str) -> None:
    table.rows[row].cells[col].text = value


def insert_table_after(paragraph: Paragraph, rows: int, cols: int) -> Table:
    tbl = OxmlElement("w:tbl")
    paragraph._p.addnext(tbl)
    table = Table(tbl, paragraph._parent)
    for _ in range(rows):
        tr = OxmlElement("w:tr")
        for _ in range(cols):
            tc = OxmlElement("w:tc")
            p = OxmlElement("w:p")
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    return table


def fill_table(table: Table, headers: list[str], data: list[list[str]]) -> None:
    for col, header in enumerate(headers):
        set_cell(table, 0, col, header)
    for row_idx, row in enumerate(data, start=1):
        for col_idx, value in enumerate(row):
            set_cell(table, row_idx, col_idx, value)


def main() -> None:
    shutil.copy2(SRC, BACKUP)

    doc = Document(str(SRC))

    start_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "IV. EXPERIMENTAL RESULTS")
    end_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "CONCLUSION AND RECOMMENDATION")

    anchor = doc.paragraphs[start_idx]
    for idx in range(end_idx - 1, start_idx, -1):
        delete_paragraph(doc.paragraphs[idx])

    sections: list[tuple[str, str]] = [
        (
            "Body Text",
            "This section reports datasets, experimental apparatus, training procedures, and quantitative findings "
            "for the TBhon multimodal tuberculosis pre-screening system. Evaluation encompassed unimodal cough and "
            "sputum classifiers, three-fold cross-validation, architectural ablation (convolutional baseline versus "
            "hybrid CNN+gradient-boosted ensemble), operating-point calibration, pre-inference cough quality control, "
            "and end-to-end API response behavior. Screening-level multimodal risk integration is described "
            "qualitatively; a dedicated held-out evaluation of the full fused triage score against microbiological "
            "ground truth was not completed and is noted as a limitation in Section IV-G.",
        ),
        ("Heading 2", "A. Datasets and Experimental Setup"),
        (
            "Body Text",
            "The evaluation employed two publicly available data sources aligned with the TBhon training pipeline.",
        ),
        (
            "Body Text",
            "Cough audio. The open Kaggle tuberculosis audio corpus provided 7,817 labeled respiratory recordings. "
            "All waveforms were converted to mono, resampled to 16 kHz, and normalized to a fixed temporal window "
            "before feature extraction. The deployed hybrid classifier used four-second clips and sixty-four Mel "
            "frequency bins; supplementary convolutional baselines additionally examined six-second clips with "
            "128 Mel bins. Labels were binary (tuberculosis-positive versus tuberculosis-negative). The publisher "
            "supplied three predefined cross-validation folds, each allocating approximately 5,211 recordings to "
            "training and 2,606 to testing. An additional 12% of the training partition was withheld as a validation "
            "set through stratified sampling (fixed random seed), preserving class prevalence for early stopping, "
            "threshold selection, and branch blending.",
        ),
        (
            "Body Text",
            "Sputum smear images. A curated set of 1,438 Ziehl–Neelsen stained microscopy fields was partitioned "
            "into 1,149 training, 142 validation, and 147 test images (approximately 80% / 10% / 10%). "
            "Acid-fast bacillus (AFB) labels were derived from bounding-box annotations by counting visible bacilli "
            "per field. Two tasks were evaluated: binary AFB-negative versus AFB-positive classification (production "
            "path) and a complementary four-class AFB load-grade formulation. Binary classification employed "
            "224 × 224-pixel inputs with a ResNet-18 backbone; the load-grade network used 160 × 160-pixel inputs.",
        ),
        (
            "Body Text",
            "Prototype acquisition. Field samples were collected using an ESP32-S3 microcontroller paired with an "
            "INMP441 I2S MEMS microphone (16 kHz audio, up to ten seconds per event, minimum three seconds in the "
            "screening workflow) and a microscope-mounted camera acquiring 640 × 480-pixel JPEG frames. Recordings "
            "and images were transmitted over Wi-Fi to a FastAPI inference server for preprocessing and classification.",
        ),
        (
            "Body Text",
            "Compute environment. Model training and offline evaluation were performed on a workstation with an "
            "Intel Core i5 processor, 8 GB RAM, and a CUDA-enabled NVIDIA GPU, using PyTorch 2.x, Librosa, and OpenCV. "
            "Production inference was deployed on a cloud-hosted Python/FastAPI service (DigitalOcean ML droplet) "
            "exposing /check-quality, /predict, and /predict-phlegm endpoints.",
        ),
        ("Heading 2", "B. Cough Audio Classification"),
        (
            "Body Text",
            "Two cough classifiers were compared on held-out test data from fold 0 (n = 2,606). The baseline was a "
            "convolutional neural network operating on log-Mel spectrograms (128 Mel bins, six-second clips), trained "
            "with AdamW (learning rate 3 × 10⁻⁴), batch size 32, up to twenty epochs, and early stopping after five "
            "epochs without validation macro-F1 improvement. Augmentation included temporal shifting, additive Gaussian "
            "noise, and time–frequency masking.",
        ),
        (
            "Body Text",
            "The production hybrid model combined a convolutional branch (sixty-four Mel bins, four-second clips, "
            "eight CNN epochs) with a gradient-boosted ensemble fed by hand-crafted acoustic descriptors (MFCCs, "
            "temporal derivatives, and summary statistics). CNN and ensemble posterior probabilities were fused through "
            "validation-tuned weighted averaging (production blend: 5% CNN / 95% GBM on fold 1). Additional augmentation "
            "simulating lossy compression and reverberation was applied during hybrid training to improve robustness to "
            "field recording conditions.",
        ),
        (
            "Body Text",
            "Table I summarizes the fold-0 ablation. The hybrid model improved macro-F1 by 17.8 percentage points "
            "over the convolutional baseline (50.6% to 68.3%), confirming that stacked spectral learning and "
            "hand-crafted features are complementary for tuberculosis-related cough discrimination under severe class "
            "imbalance.",
        ),
        ("table head", "TABLE I\nCOUGH CLASSIFIER ABLATION ON FOLD 0 (HELD-OUT TEST, n = 2,606)"),
        ("Heading 2", "C. Cross-Fold Generalization"),
        (
            "Body Text",
            "To assess stability across publisher-defined partitions, the hybrid architecture was independently trained "
            "and evaluated on all three cross-validation folds using the same preprocessing, validation split, and "
            "hyper-parameter search protocol. Table II reports held-out test metrics. Mean macro-F1 across folds was "
            "68.6% (range 68.3–69.1 pp), indicating consistent generalization despite fold-wise shifts in speaker and "
            "recording conditions. Fold 1 (production deployment) achieved the highest macro-F1 (69.1%) and was selected "
            "for integration into the mobile application.",
        ),
        ("table head", "TABLE II\nHYBRID COUGH CLASSIFIER CROSS-FOLD TEST PERFORMANCE"),
        ("Heading 2", "D. Sputum Smear Classification"),
        (
            "Body Text",
            "The binary AFB classifier employed a ResNet-18 backbone pretrained on ImageNet, fine-tuned for thirty "
            "epochs with AdamW at a learning rate of 1 × 10⁻³, batch size 32, and class-weighted cross-entropy to "
            "mitigate label imbalance. Training-time augmentation included random resized cropping, horizontal and "
            "vertical flipping, in-plane rotation, and color jitter. The decision threshold on the positive-class "
            "probability was selected on the validation set under a minimum AFB-positive sensitivity constraint of 95%, "
            "yielding a production threshold of 0.63.",
        ),
        (
            "Body Text",
            "On the held-out test partition (n = 216; 6 AFB-negative, 210 AFB-positive), the production model "
            "(run phlegm_afb_binary_20260531_133949) achieved 94.91% accuracy, 66.32% macro-F1, and 96.19% "
            "AFB-positive sensitivity (Table III). AFB-negative specificity was 50.00% (3 true negatives, 3 false "
            "positives), reflecting the screening-oriented threshold policy and the extremely small negative test "
            "cohort. These results support triage prioritization rather than standalone clinical diagnosis.",
        ),
        ("table head", "TABLE III\nPRODUCTION UNIMODAL CLASSIFIER PERFORMANCE ON HELD-OUT TEST DATA"),
        ("Heading 2", "E. Pre-Inference Quality Control"),
        (
            "Body Text",
            "Before cough classification, each recording passes a lightweight heuristic quality gate implemented in "
            "the FastAPI /check-quality endpoint. The gate rejects obvious silence, clipped replay, conversational "
            "speech, and steady background noise using RMS level, crest factor, spectral flatness, autocorrelation "
            "periodicity, dynamic range, and burst-structure heuristics calibrated on the TB-Audio corpus. Invalid "
            "clips are excluded from probability fusion rather than forced through the classifier. Functional testing "
            "with synthetic validation files confirmed correct rejection of noise, silence, and speech-like inputs while "
            "accepting transient cough bursts.",
        ),
        ("Heading 2", "F. Multimodal Screening Fusion"),
        (
            "Body Text",
            "End-to-end tuberculosis risk scoring integrates up to three independent information sources through "
            "weighted log-odds fusion: an eleven-item symptom and exposure checklist (weight 0.85), mean cough machine "
            "learning probability across quality-validated clips (weight 1.00), and sputum machine learning probability "
            "(weight 0.70). For each available modality i with probability estimate pᵢ and weight wᵢ, the fused score is",
        ),
        (
            "Body Text",
            "P_fused = σ( Σᵢ wᵢ · logit(pᵢ) / Σᵢ wᵢ )",
        ),
        (
            "Body Text",
            "where σ denotes the logistic function and logit(p) = ln(p / (1 − p)). Clinical safety floors elevate the "
            "fused probability when the checklist indicates high concern or when the sputum classifier reports a "
            "confident AFB-positive finding. The fused probability is mapped to Low Risk (< 0.38), Moderate Risk "
            "(0.38–0.62), or High Risk (≥ 0.62) for display in the mobile application. Modality reliability weights "
            "are summarized in Table IV.",
        ),
        ("table head", "TABLE IV\nMULTIMODAL FUSION MODALITY RELIABILITY WEIGHTS"),
        ("Heading 2", "G. Evaluation Metrics, System Response, and Discussion"),
        (
            "Body Text",
            "Classification performance was quantified using accuracy, precision, recall, and macro-averaged F1-score "
            "on held-out test partitions. Macro-averaging treated all classes equally and was preferred where class "
            "prevalence was skewed. A fixed random seed ensured reproducibility of validation splits and training "
            "initialization.",
        ),
        (
            "Body Text",
            "Classification response time was evaluated against the five-second per-inference target specified in the "
            "system requirements. During user acceptance testing on TBhon v1.2, four of six evaluators rated cough and "
            "sputum inference at Agree or Strongly Agree for the ≤5-second target under stable booth Wi-Fi. Occasional "
            "delays exceeding five seconds were reported on unstable networks, consistent with connectivity being the "
            "primary operational constraint rather than model compute time on the ML server.",
        ),
        (
            "Body Text",
            "Table III consolidates production unimodal performance on held-out test data. The hybrid cough classifier "
            "exhibited strong non-TB recall (88.25%), suitable for ruling out lower-risk cases during a session, while "
            "TB-positive recall remained moderate (47.63%), indicating that cough audio alone should not be treated as "
            "confirmatory. The sputum classifier prioritized AFB-positive sensitivity (96.19%) at the expense of "
            "specificity on the limited negative test set.",
        ),
        (
            "Body Text",
            "End-to-end functional testing on two mobile devices (iPhone 15 and Redmi Note 12 Pro 5G) verified the "
            "complete screening workflow: staff login, client intake, symptom checklist, cough and sputum capture, ML "
            "review, fused risk disclosure, and history retrieval. All six workflow test cases passed on both devices. "
            "Because inference is cloud-hosted, identical media inputs produced consistent API outputs; device "
            "differences were observed mainly in capture quality, upload time, and UI responsiveness.",
        ),
        (
            "Body Text",
            "Limitations. (1) Evaluation used publicly available datasets and prototype deployment conditions rather "
            "than prospective clinical studies. (2) Cough TB-class recall remains below clinical confirmation standards "
            "for standalone use. (3) Sputum specificity estimates are unstable because only six AFB-negative test "
            "images were available. (4) No dedicated held-out benchmark was completed for the full fused triage score "
            "against microbiological ground truth. (5) Response-time measurements were observational during UAT rather "
            "than instrumented server-side latency profiling. These constraints bound the present results to engineering "
            "feasibility and user acceptability rather than clinical effectiveness.",
        ),
        (
            "Body Text",
            "Despite these limitations, the experiments demonstrate that a hybrid cough classifier, a sensitivity-tuned "
            "sputum AFB classifier, and interpretable multimodal fusion can be integrated into an IoT-assisted mobile "
            "screening architecture with acceptable booth-operator usability. The multimodal design is consistent with "
            "the study hypothesis that combining respiratory audio with microscopic sputum imagery yields a more robust "
            "preliminary TB-risk indicator than either modality alone, even when unimodal metrics remain imperfect.",
        ),
    ]

    cursor = anchor
    for style, text in sections:
        cursor = insert_paragraph_after(cursor, text, style)

    # Table 0: cough ablation fold 0
    t0 = doc.tables[0]
    set_cell(t0, 0, 0, "Model")
    set_cell(t0, 0, 1, "Fold")
    set_cell(t0, 0, 2, "Test n")
    set_cell(t0, 0, 3, "Accuracy")
    set_cell(t0, 0, 4, "Precision")
    set_cell(t0, 0, 5, "Recall")
    set_cell(t0, 0, 6, "Macro-F1")
    set_cell(t0, 1, 0, "Convolutional network (baseline)")
    set_cell(t0, 1, 1, "0")
    set_cell(t0, 1, 2, "2,606")
    set_cell(t0, 1, 3, "50.6%")
    set_cell(t0, 1, 4, "60.4%")
    set_cell(t0, 1, 5, "60.7%")
    set_cell(t0, 1, 6, "50.6%")
    set_cell(t0, 2, 0, "Hybrid CNN+GBM (ablation)")
    set_cell(t0, 2, 1, "0")
    set_cell(t0, 2, 2, "2,606")
    set_cell(t0, 2, 3, "73.9%")
    set_cell(t0, 2, 4, "68.3%")
    set_cell(t0, 2, 5, "68.4%")
    set_cell(t0, 2, 6, "68.3%")

    # Expand table 0 to 5 rows for cross-fold (Table II) - easier to repurpose table 2
    # Table 1 stays fusion weights
    t1 = doc.tables[1]
    set_cell(t1, 0, 0, "Modality")
    set_cell(t1, 0, 1, "Reliability Weight")
    set_cell(t1, 1, 0, "Symptom and exposure checklist")
    set_cell(t1, 1, 1, "0.85")
    set_cell(t1, 2, 0, "Cough machine learning probability")
    set_cell(t1, 2, 1, "1.00")
    set_cell(t1, 3, 0, "Sputum machine learning probability")
    set_cell(t1, 3, 1, "0.70")

    # Table 2: production unimodal metrics (replace fabricated fusion table)
    t2 = doc.tables[2]
    set_cell(t2, 0, 0, "Production Model")
    set_cell(t2, 0, 1, "Test n")
    set_cell(t2, 0, 2, "Accuracy")
    set_cell(t2, 0, 3, "Macro Precision")
    set_cell(t2, 0, 4, "Macro Recall")
    set_cell(t2, 0, 5, "Macro-F1")
    set_cell(t2, 1, 0, "Hybrid cough CNN+GBM (fold 1)")
    set_cell(t2, 1, 1, "2,606")
    set_cell(t2, 1, 2, "75.8%")
    set_cell(t2, 1, 3, "71.7%")
    set_cell(t2, 1, 4, "67.9%")
    set_cell(t2, 1, 5, "69.1%")
    set_cell(t2, 2, 0, "ResNet-18 AFB binary (sputum)")
    set_cell(t2, 2, 1, "216")
    set_cell(t2, 2, 2, "94.9%")
    set_cell(t2, 2, 3, "62.6%")
    set_cell(t2, 2, 4, "73.0%")
    set_cell(t2, 2, 5, "66.3%")
    set_cell(t2, 3, 0, "AFB+ sensitivity (sputum)")
    set_cell(t2, 3, 1, "210 positives")
    set_cell(t2, 3, 2, "—")
    set_cell(t2, 3, 3, "98.5%")
    set_cell(t2, 3, 4, "96.2%")
    set_cell(t2, 3, 5, "97.1%")

    # Insert Table II (cross-fold) after its caption paragraph
    table_ii_caption = next(
        p for p in doc.paragraphs if p.text.strip().startswith("TABLE II")
    )
    t_cross = insert_table_after(table_ii_caption, 5, 6)
    fill_table(
        t_cross,
        ["Fold", "Test n", "Accuracy", "Macro-F1", "TB Recall", "Non-TB Recall"],
        [
            ["0", "2,606", "73.9%", "68.3%", "55.2%", "81.5%"],
            ["1 (production)", "2,606", "75.8%", "69.1%", "47.6%", "88.3%"],
            ["2", "2,605", "76.0%", "68.3%", "44.0%", "90.0%"],
            ["Mean", "—", "75.2%", "68.6%", "48.9%", "86.6%"],
        ],
    )

    doc.save(str(SRC))
    print(f"Updated {SRC}")
    print(f"Backup at {BACKUP}")


if __name__ == "__main__":
    main()
