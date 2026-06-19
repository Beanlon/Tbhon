"""Generate Appendix D — Data Augmentation Documentation (ADET format, TBhon)."""
from __future__ import annotations

import sys
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from PIL import Image, ImageDraw, ImageEnhance, ImageOps
from scipy.io import wavfile

ROOT = Path(__file__).resolve().parents[1]
DATASET = ROOT / "ml (phlegm)" / "Raw_Sputum_Microscopy_Dataset"
COUGH_SAMPLE_WAV = ROOT / "docs" / "appendix_b" / "cough" / "tb" / "wav" / "tb_sample_01.wav"
APPENDIX_D = ROOT / "docs" / "appendix_d"
FIG_DIR = ROOT / "docs" / "figures"
OUT_DOCX = ROOT / "docs" / "TBhon_Appendix_D.docx"
OUT_FALLBACK = ROOT / "docs" / "TBhon_Appendix_D_generated.docx"

# Representative AFB-positive fields from the real dataset.
SAMPLE_IMAGES = [
    DATASET / "images" / "train" / "sputum_train_0005.jpg",
    DATASET / "images" / "train" / "sputum_train_0124.png",
    DATASET / "images" / "train" / "sputum_train_0280.jpg",
]

AUG_FIGURES = {
    "rotation": APPENDIX_D / "d1_rotation_samples.png",
    "flipping": APPENDIX_D / "d2_flipping_samples.png",
    "brightness": APPENDIX_D / "d3_brightness_samples.png",
    "saturation": APPENDIX_D / "d4_saturation_samples.png",
    "gaussian_noise": APPENDIX_D / "d5_gaussian_noise_samples.png",
    "workflow": APPENDIX_D / "d6_augmentation_workflow.png",
}


def parse_yolo_boxes(label_path: Path) -> list[tuple[float, float, float, float]]:
    boxes: list[tuple[float, float, float, float]] = []
    if not label_path.is_file():
        return boxes
    with label_path.open("r", encoding="utf-8", errors="ignore") as fh:
        for line in fh:
            parts = line.strip().split()
            if len(parts) < 5:
                continue
            try:
                _, xc, yc, w, h = parts[:5]
                boxes.append((float(xc), float(yc), float(w), float(h)))
            except ValueError:
                continue
    return boxes


def load_rgb(img_path: Path) -> Image.Image:
    return Image.open(img_path).convert("RGB")


def draw_boxes(img: Image.Image, label_path: Path) -> Image.Image:
    out = img.copy()
    w, h = out.size
    draw = ImageDraw.Draw(out)
    for xc, yc, bw, bh in parse_yolo_boxes(label_path):
        x1 = (xc - bw / 2) * w
        y1 = (yc - bh / 2) * h
        x2 = (xc + bw / 2) * w
        y2 = (yc + bh / 2) * h
        draw.rectangle([x1, y1, x2, y2], outline="#E53935", width=max(2, int(min(w, h) * 0.004)))
    return out


def label_path_for(img_path: Path) -> Path:
    split = img_path.parent.name
    return DATASET / "labels" / split / f"{img_path.stem}.txt"


def rotate(img: Image.Image, degrees: float) -> Image.Image:
    return img.rotate(degrees, resample=Image.Resampling.BICUBIC, expand=False, fillcolor=(0, 0, 0))


def flip_h(img: Image.Image) -> Image.Image:
    return ImageOps.mirror(img)


def flip_v(img: Image.Image) -> Image.Image:
    return ImageOps.flip(img)


def adjust_brightness(img: Image.Image, factor: float) -> Image.Image:
    return ImageEnhance.Brightness(img).enhance(factor)


def adjust_saturation(img: Image.Image, factor: float) -> Image.Image:
    return ImageEnhance.Color(img).enhance(factor)


def pil_to_array(img: Image.Image) -> np.ndarray:
    return np.asarray(img)


def save_comparison_panel(
    out_path: Path,
    title: str,
    rows: list[tuple[str, list[tuple[str, Image.Image]]]],
) -> None:
    """rows: [(row_label, [(caption, image), ...]), ...]"""
    n_rows = len(rows)
    n_cols = max(len(cells) for _, cells in rows)
    fig, axes = plt.subplots(n_rows, n_cols, figsize=(3.0 * n_cols, 3.0 * n_rows))
    if n_rows == 1 and n_cols == 1:
        axes = np.array([[axes]])
    elif n_rows == 1:
        axes = np.array([axes])
    elif n_cols == 1:
        axes = axes.reshape(n_rows, 1)

    fig.patch.set_facecolor("white")
    fig.suptitle(title, fontsize=13, fontweight="bold", y=0.98)

    for r, (_, cells) in enumerate(rows):
        for c in range(n_cols):
            ax = axes[r, c]
            ax.axis("off")
            if c < len(cells):
                caption, img = cells[c]
                ax.imshow(pil_to_array(img))
                ax.set_title(caption, fontsize=10, fontweight="bold", pad=6)
            else:
                ax.set_visible(False)

    plt.tight_layout(rect=[0, 0, 1, 0.96])
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def resolve_samples() -> list[tuple[Image.Image, Path]]:
    samples: list[tuple[Image.Image, Path]] = []
    for img_path in SAMPLE_IMAGES:
        if not img_path.is_file():
            continue
        lbl = label_path_for(img_path)
        samples.append((draw_boxes(load_rgb(img_path), lbl), img_path))
    if not samples:
        raise FileNotFoundError("No sample sputum images found for Appendix D.")
    return samples


def generate_d1_rotation(samples: list[tuple[Image.Image, Path]]) -> None:
    rows: list[tuple[str, list[tuple[str, Image.Image]]]] = []
    for img, _ in samples[:3]:
        rows.append(
            (
                "",
                [
                    ("Original", img),
                    ("Rotated +15°", rotate(img, 15)),
                    ("Rotated −15°", rotate(img, -15)),
                ],
            )
        )
    save_comparison_panel(
        AUG_FIGURES["rotation"],
        "D.1 Rotation Samples — RandomRotation(±15°) applied to sputum smear fields",
        rows,
    )


def generate_d2_flipping(samples: list[tuple[Image.Image, Path]]) -> None:
    img = samples[0][0]
    save_comparison_panel(
        AUG_FIGURES["flipping"],
        "D.2 Flipping Samples — RandomHorizontalFlip / RandomVerticalFlip (p=0.5)",
        [
            (
                "",
                [
                    ("Original", img),
                    ("Horizontal flip", flip_h(img)),
                    ("Vertical flip", flip_v(img)),
                ],
            )
        ],
    )


def generate_d3_brightness(samples: list[tuple[Image.Image, Path]]) -> None:
    img = samples[1][0] if len(samples) > 1 else samples[0][0]
    save_comparison_panel(
        AUG_FIGURES["brightness"],
        "D.3 Brightness Adjustment Samples — ColorJitter(brightness=0.15)",
        [
            (
                "",
                [
                    ("Original", img),
                    ("Brightness +15%", adjust_brightness(img, 1.15)),
                    ("Brightness −15%", adjust_brightness(img, 0.85)),
                ],
            )
        ],
    )


def generate_d4_saturation(samples: list[tuple[Image.Image, Path]]) -> None:
    img = samples[2][0] if len(samples) > 2 else samples[0][0]
    save_comparison_panel(
        AUG_FIGURES["saturation"],
        "D.4 Saturation Adjustment Samples — ColorJitter(saturation=0.15)",
        [
            (
                "",
                [
                    ("Original", img),
                    ("Saturation +15%", adjust_saturation(img, 1.15)),
                    ("Saturation −15%", adjust_saturation(img, 0.85)),
                ],
            )
        ],
    )


def load_wav_mono(path: Path) -> tuple[np.ndarray, int]:
    sr, data = wavfile.read(path)
    arr = np.asarray(data)
    if arr.ndim == 2:
        arr = arr.mean(axis=1)
    if np.issubdtype(arr.dtype, np.integer):
        max_val = float(np.iinfo(arr.dtype).max)
        arr = arr.astype(np.float32) / max_val
    else:
        arr = arr.astype(np.float32)
    return arr, int(sr)


def clip_or_pad(wav: np.ndarray, target_len: int) -> np.ndarray:
    if len(wav) >= target_len:
        start = max(0, (len(wav) - target_len) // 2)
        return wav[start : start + target_len]
    out = np.zeros(target_len, dtype=np.float32)
    out[: len(wav)] = wav
    return out


def add_waveform_gaussian_noise(wav: np.ndarray, noise_std: float, seed: int) -> np.ndarray:
    rng = np.random.default_rng(seed)
    return wav + rng.normal(0.0, noise_std, wav.shape).astype(np.float32)


def generate_d5_gaussian_noise() -> None:
    """Cough pipeline: x + randn * noise_std (train_tb_cough_hybrid.py, σ=0.008)."""
    if not COUGH_SAMPLE_WAV.is_file():
        raise FileNotFoundError(f"Cough sample not found for D.5: {COUGH_SAMPLE_WAV}")

    sample_rate = 16000
    clip_seconds = 4.0
    noise_std = 0.008
    wav, sr = load_wav_mono(COUGH_SAMPLE_WAV)
    if sr != sample_rate:
        duration = len(wav) / sr
        target_len = max(1, int(round(duration * sample_rate)))
        x_old = np.linspace(0.0, 1.0, num=len(wav), endpoint=False)
        x_new = np.linspace(0.0, 1.0, num=target_len, endpoint=False)
        wav = np.interp(x_new, x_old, wav).astype(np.float32)
    target_len = int(sample_rate * clip_seconds)
    wav = clip_or_pad(wav, target_len)
    times = np.arange(len(wav)) / sample_rate

    panels = [
        ("Original waveform", wav),
        (f"noise_std={noise_std}", add_waveform_gaussian_noise(wav, noise_std, seed=42)),
        (f"noise_std={noise_std * 2}", add_waveform_gaussian_noise(wav, noise_std * 2, seed=43)),
    ]

    fig, axes = plt.subplots(1, 3, figsize=(12, 3.2), facecolor="white")
    fig.suptitle(
        "D.5 Gaussian Noise Samples — additive waveform noise (train_tb_cough_hybrid.py)",
        fontsize=12,
        fontweight="bold",
        y=1.02,
    )
    for ax, (caption, series) in zip(axes, panels):
        ax.plot(times, series, color="#6A1B9A", linewidth=0.7)
        ax.set_xlim(0, clip_seconds)
        ax.set_title(caption, fontsize=10, fontweight="bold")
        ax.set_xlabel("Time (s)", fontsize=9)
        ax.set_ylabel("Amplitude", fontsize=9)
        ax.grid(True, alpha=0.25, linewidth=0.5)
    plt.tight_layout()
    AUG_FIGURES["gaussian_noise"].parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(AUG_FIGURES["gaussian_noise"], dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def _box(ax, x, y, w, h, text, fc="#E3F2FD", ec="#1565C0", fontsize=8.5):
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.02",
        linewidth=1.4,
        edgecolor=ec,
        facecolor=fc,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, wrap=True)


def _arrow(ax, x1, y1, x2, y2):
    ax.add_patch(
        FancyArrowPatch(
            (x1, y1),
            (x2, y2),
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=1.4,
            color="#37474F",
        )
    )


def generate_d6_workflow() -> None:
    fig = plt.figure(figsize=(14, 8.5), facecolor="white")
    ax = fig.add_axes([0.04, 0.08, 0.92, 0.82])
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.96, "D.6 Augmentation Workflow — TBhon Multimodal Training Pipeline", ha="center", fontsize=13, fontweight="bold")

    # Sputum branch (top)
    ax.text(0.5, 0.88, "Sputum microscopy (ResNet18 CNN)", ha="center", fontsize=11, fontweight="bold", color="#1565C0")
    sx = [0.03, 0.19, 0.35, 0.51, 0.67, 0.83]
    sy = 0.72
    bw, bh = 0.13, 0.12
    sputum_steps = [
        "Raw smear\nimage",
        "RandomResized\nCrop 224",
        "Rotation\n±15°",
        "H/V flip\n(p=0.5)",
        "ColorJitter\nB/S/C/H",
        "Normalize\n→ CNN",
    ]
    colors = ["#E8F5E9", "#E3F2FD", "#FFF3E0", "#FFF8E1", "#FFE0B2", "#E1F5FE"]
    edges = ["#2E7D32", "#1565C0", "#EF6C00", "#F9A825", "#E65100", "#0277BD"]
    for i, (x, lab, fc, ec) in enumerate(zip(sx, sputum_steps, colors, edges)):
        _box(ax, x, sy, bw, bh, lab, fc=fc, ec=ec, fontsize=8)
        if i < len(sx) - 1:
            _arrow(ax, x + bw + 0.005, sy + bh / 2, sx[i + 1] - 0.005, sy + bh / 2)

    # Cough branch (bottom)
    ax.text(0.5, 0.58, "Cough audio (Hybrid CNN + GBM)", ha="center", fontsize=11, fontweight="bold", color="#6A1B9A")
    cy = 0.42
    cough_steps = [
        "Raw .wav\n16 kHz",
        "Clip / pad\n4 s",
        "Time shift\n±15%",
        "Gaussian\nnoise σ=0.008",
        "Codec / reverb\n(35% / 20%)",
        "Mel + SpecAug\n→ CNN",
    ]
    cough_colors = ["#F3E5F5", "#EDE7F6", "#E8EAF6", "#E1F5FE", "#FFF3E0", "#FCE4EC"]
    cough_edges = ["#6A1B9A", "#4527A0", "#283593", "#0277BD", "#EF6C00", "#AD1457"]
    for i, (x, lab, fc, ec) in enumerate(zip(sx, cough_steps, cough_colors, cough_edges)):
        _box(ax, x, cy, bw, bh, lab, fc=fc, ec=ec, fontsize=8)
        if i < len(sx) - 1:
            _arrow(ax, x + bw + 0.005, cy + bh / 2, sx[i + 1] - 0.005, cy + bh / 2)

    ax.text(
        0.5,
        0.18,
        "Purpose: increase dataset diversity, reduce overfitting, and improve robustness under varying "
        "microscopy staining, capture orientation, lighting, and mobile/IoT audio quality.",
        ha="center",
        fontsize=9.5,
        color="#455A64",
        wrap=True,
    )
    ax.text(
        0.5,
        0.08,
        "Source code: ml (phlegm)/train_phlegm_cnn.py  ·  ml/train_tb_cough_hybrid.py",
        ha="center",
        fontsize=9,
        style="italic",
        color="#666666",
    )

    AUG_FIGURES["workflow"].parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(AUG_FIGURES["workflow"], dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int = 2) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, *, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_figure(doc: Document, path: Path, caption: str, width_in: float = 6.2) -> None:
    if path.is_file():
        doc.add_picture(str(path), width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"[Insert figure: {path.name}]", italic=True)
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_word_doc() -> Document:
    doc = Document()
    style_doc(doc)

    title = doc.add_heading("APPENDIX D — DATA AUGMENTATION DOCUMENTATION", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "This appendix demonstrates preprocessing and augmentation procedures applied to the TBhon "
        "training datasets. Sputum smear examples (D.1–D.4) are taken from Raw_Sputum_Microscopy_Dataset "
        "with YOLO AFB bounding boxes drawn for reference. Cough waveform noise (D.5) uses the same "
        "additive Gaussian transform as train_tb_cough_hybrid.py. All parameters below are taken directly "
        "from ml (phlegm)/train_phlegm_cnn.py and ml/train_tb_cough_hybrid.py.",
    )

    sections = [
        (
            "D.1  Rotation Samples",
            "RandomRotation(±15°) in train_phlegm_cnn.py simulates slight orientation differences during "
            "microscopy capture and mobile camera alignment. Three representative AFB-positive fields show "
            "original and ±15° variants.",
            AUG_FIGURES["rotation"],
            "Figure D.1. Sputum rotation augmentation samples (±15°).",
        ),
        (
            "D.2  Flipping Samples",
            "RandomHorizontalFlip and RandomVerticalFlip (p=0.5) increase spatial invariance. Microscopy "
            "fields have no fixed orientation, so flips are valid label-preserving transforms.",
            AUG_FIGURES["flipping"],
            "Figure D.2. Sputum horizontal and vertical flip samples.",
        ),
        (
            "D.3  Brightness Adjustment Samples",
            "ColorJitter(brightness=0.15, contrast=0.15, saturation=0.15, hue=0.03) in train_phlegm_cnn.py "
            "adjusts exposure to mimic variable microscope illumination and mobile camera auto-exposure. "
            "Brightness samples show the ±15% range.",
            AUG_FIGURES["brightness"],
            "Figure D.3. Sputum brightness adjustment samples (±15%).",
        ),
        (
            "D.4  Saturation Adjustment Samples",
            "The same ColorJitter block applies saturation=0.15 (and contrast=0.15, hue=0.03) to simulate "
            "differences in Ziehl–Neelsen staining intensity across preparation batches.",
            AUG_FIGURES["saturation"],
            "Figure D.4. Sputum saturation adjustment samples (±15%).",
        ),
        (
            "D.5  Gaussian Noise Samples",
            "In train_tb_cough_hybrid.py, _augment_waveform() adds Gaussian noise to the 16 kHz cough "
            "waveform before log-mel conversion: x = x + randn * noise_std with noise_std=0.008. "
            "Sputum CNN training does not apply pixel-level noise; only ColorJitter and geometric transforms.",
            AUG_FIGURES["gaussian_noise"],
            "Figure D.5. Cough waveform Gaussian noise samples (noise_std=0.008, production hybrid config).",
        ),
        (
            "D.6  Augmentation Workflow",
            "End-to-end augmentation workflow for both sputum CNN and cough hybrid model training pipelines.",
            AUG_FIGURES["workflow"],
            "Figure D.6. TBhon multimodal data augmentation workflow.",
        ),
    ]

    for heading, body, fig_path, caption in sections:
        add_heading(doc, heading)
        add_para(doc, body)
        add_figure(doc, fig_path, caption)

    add_para(
        doc,
        "Sputum training transforms (train_phlegm_cnn.py): RandomResizedCrop(224, scale=(0.85, 1.0), "
        "ratio=(0.9, 1.1)), RandomHorizontalFlip(0.5), RandomVerticalFlip(0.5), RandomRotation(15), "
        "ColorJitter(brightness=0.15, contrast=0.15, saturation=0.15, hue=0.03), ImageNet normalization.",
        italic=True,
    )
    add_para(
        doc,
        "Cough hybrid transforms (train_tb_cough_hybrid.py): time_shift_max=0.15, noise_std=0.008, "
        "codec_aug_prob=0.35, reverb_aug_prob=0.20, SpecAugment (2 frequency masks, param=16; "
        "2 time masks, param=24). Augmentation is training-only; validation and test use deterministic preprocessing.",
        italic=True,
    )

    return doc


def save_doc(doc: Document) -> Path:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(OUT_DOCX))
        return OUT_DOCX
    except PermissionError:
        doc.save(str(OUT_FALLBACK))
        print(f"[warn] Could not overwrite {OUT_DOCX}. Saved to {OUT_FALLBACK}")
        return OUT_FALLBACK


def main() -> int:
    if not DATASET.is_dir():
        print(f"Dataset not found: {DATASET}", file=sys.stderr)
        return 1

    samples = resolve_samples()
    generate_d1_rotation(samples)
    generate_d2_flipping(samples)
    generate_d3_brightness(samples)
    generate_d4_saturation(samples)
    generate_d5_gaussian_noise()
    generate_d6_workflow()

    out = save_doc(build_word_doc())

    print("Generated:")
    print(f"  {out}")
    for key, path in AUG_FIGURES.items():
        print(f"  {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
