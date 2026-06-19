"""Generate Appendix C — Dataset Annotation Documentation (ADET format, TBhon)."""
from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "docs" / "figures"
APPENDIX_C = ROOT / "docs" / "appendix_c"
OUT_DOCX = ROOT / "docs" / "TBhon_Appendix_C.docx"
OUT_FALLBACK = ROOT / "docs" / "TBhon_Appendix_C_generated.docx"

SPUTUM_DATASET = ROOT / "ml (phlegm)" / "Raw_Sputum_Microscopy_Dataset"
DATA_YAML = ROOT / "ml (phlegm)" / "data.yaml"
COUGH_METRICS = ROOT / "ml" / "runs" / "20260531_014419" / "metrics.json"
COUGH_CONFIG = ROOT / "ml" / "runs" / "20260531_014419" / "config.json"
SPUTUM_METRICS = ROOT / "ml (phlegm)" / "runs" / "phlegm_afb_binary_20260531_133949" / "metrics.json"
SPUTUM_CONFIG = ROOT / "ml (phlegm)" / "runs" / "phlegm_afb_binary_20260531_133949" / "config.json"
PRODUCTION = ROOT / "ml" / "production_model.json"

EXAMPLE_LABEL_SRC = SPUTUM_DATASET / "labels" / "train" / "sputum_train_0005.txt"
EXAMPLE_LABEL_DST = APPENDIX_C / "examples" / "sputum_train_0005.txt"
EXAMPLE_IMAGE = SPUTUM_DATASET / "images" / "train" / "sputum_train_0005.jpg"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def pct(value: float, digits: int = 2) -> str:
    return f"{100.0 * value:.{digits}f}%"


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_figure(doc: Document, path: Path, width_in: float = 5.8) -> bool:
    if not path.is_file():
        add_para(doc, f"[Insert figure: {path.name}]", italic=True)
        return False
    doc.add_picture(str(path), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    for line in text.strip().splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def export_label_example() -> None:
    APPENDIX_C.mkdir(parents=True, exist_ok=True)
    (APPENDIX_C / "examples").mkdir(parents=True, exist_ok=True)
    if EXAMPLE_LABEL_SRC.is_file():
        shutil.copy2(EXAMPLE_LABEL_SRC, EXAMPLE_LABEL_DST)


def build_appendix_c() -> Document:
    export_label_example()

    cough_metrics = load_json(COUGH_METRICS) if COUGH_METRICS.is_file() else {}
    cough_config = load_json(COUGH_CONFIG) if COUGH_CONFIG.is_file() else {}
    sputum_metrics = load_json(SPUTUM_METRICS) if SPUTUM_METRICS.is_file() else {}
    sputum_config = load_json(SPUTUM_CONFIG) if SPUTUM_CONFIG.is_file() else {}
    production = load_json(PRODUCTION) if PRODUCTION.is_file() else {}

    doc = Document()
    style_doc(doc)

    title = doc.add_heading("APPENDIX C — DATASET ANNOTATION DOCUMENTATION", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "This appendix presents evidence of dataset annotation, labeling procedures, training "
        "environment configuration, and model validation outputs for the TBhon multimodal TB "
        "pre-screening system. Sputum smear microscopy images use YOLO-format bounding-box "
        "annotations for acid-fast bacilli (AFB). Cough audio recordings use participant-level "
        "tb_status labels from the CODA TB Kaggle corpus, converted to log-mel spectrograms for CNN training.",
    )

    # --- C.1 ---
    add_heading(doc, "C.1  Annotation Workflow", level=2)
    add_para(
        doc,
        "Sputum annotation workflow: each visible AFB rod in a microscopy field was labeled with a "
        "normalized YOLO bounding box (class 0). Per-image AFB counts were derived from label files "
        "and mapped to load grades (none, low, moderate, high) for CNN binary classification training.",
    )
    add_para(
        doc,
        "Cough audio labeling workflow: recordings in the Kaggle ruchikashirsath/tb-audio dataset "
        "include tb_status metadata (0 = No-TB, 1 = TB-positive). Audio was resampled to 16 kHz, "
        "clipped or padded to a fixed duration, and transformed into log-mel spectrograms for hybrid "
        "CNN + gradient-boosted classifier training.",
    )
    if add_figure(doc, FIG / "figure_5_3_sputum_annotation_workflow.png", 6.2):
        add_caption(doc, "Figure C.1. Sputum dataset annotation and AFB load grading workflow.")
    if add_figure(doc, FIG / "cough_log_mel_spectrogram_example.png", 5.5):
        add_caption(doc, "Figure C.2. Cough audio log-mel spectrogram feature representation.")

    # --- C.2 ---
    add_heading(doc, "C.2  Bounding Box and Labeling Samples", level=2)
    add_para(
        doc,
        "Representative annotated sputum smear fields showing raw microscopy images, YOLO bounding "
        "boxes around individual AFB rods, and load-grade examples used during dataset preparation.",
    )
    if add_figure(doc, FIG / "sputum_annotated_sample.png", 5.8):
        add_caption(doc, "Figure C.3. Raw vs. annotated sputum smear sample (16 AFB → high load grade).")
    if add_figure(doc, FIG / "sputum_load_grade_examples.png", 6.0):
        add_caption(doc, "Figure C.4. AFB load grade examples across none, low, moderate, and high counts.")

    # --- C.3 ---
    add_heading(doc, "C.3  Label Categories and Export Format", level=2)
    add_para(
        doc,
        "Sputum YOLO labels use a single object class. Each line in a .txt label file stores "
        "class_id, x_center, y_center, width, and height (all normalized to 0–1).",
    )
    add_table(
        doc,
        ["Class ID", "Label", "Description"],
        [
            ["0", "AFB", "Acid-fast bacillus (Mycobacterium tuberculosis rod)"],
        ],
    )
    add_para(doc, "YOLO dataset configuration (data.yaml):", bold=True)
    if DATA_YAML.is_file():
        add_code_block(doc, DATA_YAML.read_text(encoding="utf-8"))
    add_para(doc, "Example label file (sputum_train_0005.txt — 16 AFB boxes):", bold=True)
    if EXAMPLE_LABEL_DST.is_file():
        preview = EXAMPLE_LABEL_DST.read_text(encoding="utf-8").strip().splitlines()[:5]
        add_code_block(doc, "\n".join(preview) + "\n... (16 lines total)")
    add_para(doc, "Cough audio class labels (CODA TB metadata):", bold=True)
    add_table(
        doc,
        ["tb_status", "Class Name", "Description"],
        [
            ["0", "No-TB", "Cough recording from participant without microbiologically confirmed TB"],
            ["1", "TB-Positive", "Cough recording from participant with confirmed TB status"],
        ],
    )
    add_para(
        doc,
        "Split CSV files (X_train_Fold_N.csv, X_test_Fold_N.csv) map filename → tb_status for "
        "reproducible cross-validation during cough model training.",
    )

    # --- C.4 ---
    add_heading(doc, "C.4  Annotation Validation and Training Environment", level=2)
    add_para(
        doc,
        "Dataset integrity was validated by checking image–label filename pairs across train, val, "
        "and test splits. Sputum labels were verified for normalized coordinate bounds and non-empty "
        "AFB-positive fields. Cough filenames in split CSVs were indexed against raw_data/*.wav.",
    )
    add_para(doc, "Table C.1. Sputum CNN training configuration (ResNet18 binary classifier)", bold=True)
    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Framework", "PyTorch / torchvision ResNet18"],
            ["Task", sputum_config.get("task", "binary")],
            ["Image size", str(sputum_config.get("img_size", 224))],
            ["Epochs", str(sputum_config.get("epochs", 30))],
            ["Batch size", str(sputum_config.get("batch_size", 32))],
            ["Augmentation", str(sputum_config.get("augment", True))],
            ["Dataset", "Raw_Sputum_Microscopy_Dataset"],
        ],
    )
    add_para(doc, "Table C.2. Cough hybrid CNN training configuration", bold=True)
    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Framework", "PyTorch CNN + gradient-boosted features"],
            ["Dataset", cough_config.get("dataset_slug", "ruchikashirsath/tb-audio")],
            ["Sample rate", f"{cough_config.get('sample_rate', 16000)} Hz"],
            ["Clip duration", f"{cough_config.get('clip_seconds', 4.0)} s"],
            ["Mel bins", str(cough_config.get("n_mels", 64))],
            ["CNN epochs", str(cough_config.get("cnn_epochs", 8))],
            ["Cross-validation fold", str(cough_config.get("fold", 1))],
        ],
    )
    if add_figure(doc, FIG / "figure_5_5_pytorch_fastapi_deployment_workflow.png", 6.2):
        add_caption(doc, "Figure C.5. PyTorch model deployment workflow (FastAPI ML droplet).")
    if add_figure(doc, FIG / "ml_droplet_architecture.png", 5.8):
        add_caption(doc, "Figure C.6. Cloud ML inference architecture.")

    # --- C.5 ---
    add_heading(doc, "C.5  Training Progress Evidence", level=2)
    add_para(
        doc,
        "Training loss and validation metric curves document model learning progress for both "
        "modalities across epochs.",
    )
    if add_figure(doc, FIG / "cough_metrics_compilation_8_epochs.png", 6.0):
        add_caption(doc, "Figure C.7. Cough hybrid model metrics compilation (8 CNN epochs).")
    if add_figure(doc, FIG / "sputum_metrics_compilation_30_epochs.png", 6.0):
        add_caption(doc, "Figure C.8. Sputum ResNet18 binary classifier metrics compilation (30 epochs).")
    if add_figure(doc, FIG / "cough_train_loss.png", 5.0):
        add_caption(doc, "Figure C.9. Cough model training loss curve.")
    if add_figure(doc, FIG / "sputum_train_loss.png", 5.0):
        add_caption(doc, "Figure C.10. Sputum model training loss curve.")

    # --- C.6 ---
    add_heading(doc, "C.6  Model Validation Evidence", level=2)
    add_para(
        doc,
        "Validation F1 and accuracy curves on held-out splits demonstrate generalization during "
        "training. Sputum validation reflects binary AFB presence; cough validation tracks macro F1 "
        "and per-class TB / No-TB performance.",
    )
    if add_figure(doc, FIG / "cough_val_f1_macro.png", 4.8):
        add_caption(doc, "Figure C.11. Cough validation macro F1.")
    if add_figure(doc, FIG / "sputum_val_f1_macro.png", 4.8):
        add_caption(doc, "Figure C.12. Sputum validation macro F1.")
    if add_figure(doc, FIG / "cough_val_f1_tb.png", 4.5):
        add_caption(doc, "Figure C.13. Cough validation F1 — TB class.")
    if add_figure(doc, FIG / "cough_val_f1_no_tb.png", 4.5):
        add_caption(doc, "Figure C.14. Cough validation F1 — No-TB class.")

    # --- C.7 ---
    add_heading(doc, "C.7  Confusion Matrix and Evaluation Metrics", level=2)
    add_para(
        doc,
        "Final test-set confusion matrices and classification metrics for the production-promoted "
        "model runs used in TBhon inference.",
    )
    cm_cough = cough_metrics.get("confusion_matrix", [[], []])
    add_table(
        doc,
        ["Cough Model (Hybrid CNN) — Test Metrics", "Value"],
        [
            ["Test accuracy", pct(cough_metrics.get("test_accuracy", 0))],
            ["Macro F1", f"{cough_metrics.get('best_f1_macro', 0):.4f}"],
            ["Decision threshold", str(cough_metrics.get("decision_threshold", "—"))],
            ["Confusion matrix [[TN,FP],[FN,TP]]", str(cm_cough)],
        ],
    )
    cm_sputum = sputum_metrics.get("confusion_matrix", [[], []])
    add_table(
        doc,
        ["Sputum Model (ResNet18 Binary) — Test Metrics", "Value"],
        [
            ["Test accuracy", pct(sputum_metrics.get("test_acc", 0))],
            ["Macro F1", f"{sputum_metrics.get('test_macro_f1', 0):.4f}"],
            ["Sensitivity", pct(sputum_metrics.get("sensitivity", 0))],
            ["Specificity", pct(sputum_metrics.get("specificity", 0))],
            ["Decision threshold", str(sputum_metrics.get("decision_threshold", "—"))],
            ["Confusion matrix [[TN,FP],[FN,TP]]", str(cm_sputum)],
        ],
    )
    if add_figure(doc, FIG / "cough_confusion_matrix.png", 4.5):
        add_caption(doc, "Figure C.15. Cough model confusion matrix.")
    if add_figure(doc, FIG / "sputum_confusion_matrix.png", 4.5):
        add_caption(doc, "Figure C.16. Sputum model confusion matrix.")
    if add_figure(doc, FIG / "figure_5_4_cough_architecture_and_confusion_matrix.png", 6.0):
        add_caption(doc, "Figure C.17. Cough CNN architecture and evaluation summary.")

    # --- C.8 ---
    add_heading(doc, "C.8  Model Export and Deployment Preparation", level=2)
    add_para(
        doc,
        "Trained model weights were exported from PyTorch training runs and registered in "
        "ml/production_model.json for FastAPI inference deployment on the ML droplet. The cough "
        "hybrid bundle requires both model.pt and hybrid_bundle.pkl; the sputum classifier uses "
        "model_best.pt from the phlegm training run.",
    )
    add_table(
        doc,
        ["Deployment Artifact", "Value"],
        [
            ["Production run ID", production.get("run_id", "20260531_014419")],
            ["Cough model type", production.get("model_type", "hybrid_cnn")],
            ["Cough model path", production.get("model_path", "runs/20260531_014419/model.pt")],
            ["Cough test accuracy", pct(production.get("test_accuracy", cough_metrics.get("test_accuracy", 0)))],
            ["Cough macro F1", f"{production.get('best_f1_macro', cough_metrics.get('best_f1_macro', 0)):.4f}"],
            ["Inference API", "ml/infer_api.py — /predict/cough, /predict/sputum"],
            ["Sputum weights", "ml (phlegm)/runs/phlegm_afb_binary_20260531_133949/model_best.pt"],
        ],
    )
    add_para(
        doc,
        "Note: All models are deployed for triage support only and are not certified diagnostic devices. "
        "Results require staff review and follow-up microbiological confirmation per program protocol.",
        italic=True,
    )

    return doc


def save_doc(doc: Document) -> Path:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(OUT_DOCX))
        return OUT_DOCX
    except PermissionError:
        doc.save(str(OUT_FALLBACK))
        print(f"[warn] Could not overwrite {OUT_DOCX} (file may be open). Saved to {OUT_FALLBACK}")
        return OUT_FALLBACK


def main() -> int:
    missing = [
        p
        for p in (
            FIG / "figure_5_3_sputum_annotation_workflow.png",
            COUGH_METRICS,
            SPUTUM_METRICS,
        )
        if not p.is_file()
    ]
    if missing:
        print("Warning: missing assets (placeholders will be inserted):", file=sys.stderr)
        for p in missing:
            print(f"  {p}", file=sys.stderr)

    doc = build_appendix_c()
    out = save_doc(doc)
    print("Generated:")
    print(f"  {out}")
    print(f"  {APPENDIX_C / 'examples'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
