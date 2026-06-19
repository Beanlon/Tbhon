"""Generate TBhon Chapter 7 — Conclusions and Recommendations (Word)."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "TBhon_Chapter_7.docx"
OUT_FALLBACK = ROOT / "docs" / "TBhon_Chapter_7_generated.docx"

COUGH_METRICS = ROOT / "ml" / "runs" / "20260531_014419" / "metrics.json"
SPUTUM_METRICS = ROOT / "ml (phlegm)" / "runs" / "phlegm_afb_binary_20260531_133949" / "metrics.json"
UAT_SUMMARY = ROOT / "docs" / "uat_summary.json"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_recommendation(doc: Document, lead: str, body: str) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(lead + " ")
    r1.bold = True
    r2 = p.add_run(body)
    for r in (r1, r2):
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)


def pct(value: float, digits: int = 2) -> str:
    return f"{100.0 * value:.{digits}f}%"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def build() -> None:
    cough = load_json(COUGH_METRICS)
    sputum = load_json(SPUTUM_METRICS)
    uat = load_json(UAT_SUMMARY) if UAT_SUMMARY.is_file() else None
    uat_n = uat["n"] if uat else 6
    grand_mean = uat["grand_mean"] if uat else 4.53
    grand_interp = uat["grand_mean_interp"] if uat else "SA"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    add_heading(doc, "CHAPTER VII", 1)
    add_heading(doc, "CONCLUSIONS AND RECOMMENDATIONS", 1)

    add_para(
        doc,
        "This chapter presents the conclusions drawn from the development, testing, evaluation, and user acceptance "
        "procedures conducted on TBhon version 1.2 (Prototype / Field Evaluation Build), together with recommendations "
        "for future research, deployment refinement, and clinical integration. The conclusions are based on the "
        "achievement of the eight specific objectives outlined in Section 1.3, the machine learning performance results "
        "in Chapter VI, functional and inference test cases, and ISO/IEC 25010–aligned user acceptance testing with "
        "six purposively selected professional evaluators.",
    )

    # --- 7.1 Conclusions ---
    add_heading(doc, "7.1 Conclusions", 2)

    add_para(
        doc,
        "The study successfully developed TBhon, an IoT-enabled mobile-based multimodal tuberculosis pre-screening and "
        "triage support system integrating an Expo/React Native mobile application, a Node.js/Express backend with "
        "Prisma ORM and DigitalOcean Managed MySQL, a Python/FastAPI machine learning inference service with PyTorch "
        "models, optional ESP32-S3 IoT acquisition hardware, and weighted multimodal risk fusion for Low, Moderate, "
        "and High triage stratification. The integration of artificial intelligence, mobile computing, cloud "
        "deployment, and IoT-assisted data capture enabled the proposed system to support staff-guided screening "
        "workflows involving client intake, an eleven-item symptom checklist, cough audio capture, sputum smear "
        "imaging, cloud ML inference, staff review, fused risk disclosure, screening history, and patient QR result "
        "claim within community and booth-oriented TB pre-screening environments.",
    )

    add_para(
        doc,
        "The first specific objective—development of IoT acquisition hardware for respiratory audio and microscopic "
        "sputum imaging—was substantially achieved. The researchers designed and implemented an ESP32-S3 module with "
        "INMP441 cough audio capture and OV5640 sputum imaging, authenticated REST upload to backend /iot endpoints, "
        "and mobile-side IoT health verification. Bench testing confirmed device-triggered media upload and session "
        "linkage. Full in-application BLE Wi-Fi provisioning requires a native deployment build and remains a "
        "refinement item rather than a blocker to the core screening workflow.",
    )

    add_para(
        doc,
        "The second specific objective—data curation and preprocessing—was achieved. Publicly available cough audio "
        "and YOLO-annotated sputum microscopy datasets were partitioned into training, validation, and test subsets. "
        "Cough samples were resampled to 16 kHz, converted to log-Mel spectrograms, and supplemented with hand-crafted "
        "acoustic features for the gradient boosting branch. Sputum images were annotated, converted to binary "
        "acid-fast bacilli labels, resized to 224×224 pixels, and augmented to improve model robustness.",
    )

    add_para(
        doc,
        f"The third specific objective—machine learning model development—was achieved. The hybrid Mel-spectrogram "
        f"CNN + Gradient Boosting Machine cough classifier (production run 20260531_014419) attained "
        f"{pct(cough['test_accuracy'])} test accuracy and {pct(cough['best_f1_macro'])} macro F1 on a held-out test "
        f"fold (n = 2,606), with strong non-TB recall (88.25%) suitable for ruling out lower-risk cases during triage. "
        f"The ResNet18 binary sputum AFB classifier (production run phlegm_afb_binary_20260531_133949) attained "
        f"{pct(sputum['test_acc'])} test accuracy and {pct(sputum['sensitivity'])} AFB-positive sensitivity on a "
        f"held-out test partition (n = 216), reflecting a screening-oriented threshold policy that prioritizes "
        f"sensitivity over specificity. Both models were deployed through FastAPI endpoints with quality gating for "
        f"invalid cough audio and unusable smear images.",
    )

    add_para(
        doc,
        "The fourth specific objective—multimodal risk scoring—was achieved. The system combined checklist probability, "
        "mean cough ML probability across quality-validated clips, and sputum ML probability through weighted log-odds "
        "fusion (weights 0.85, 1.00, and 0.70) with clinical safety floors for high-concern symptoms and confident "
        "AFB-positive findings. The fused output produced interpretable Low, Moderate, and High triage bands with "
        "disclaimers and referral-oriented guidance, extending practical utility beyond isolated modality scores.",
    )

    add_para(
        doc,
        "The fifth specific objective—communication architecture—was achieved. Two DigitalOcean cloud droplets "
        "hosted the backend (PM2, port 4000) and ML API (systemd, port 8000), both exposed through Cloudflare HTTPS "
        "tunnels to support reliable mobile uploads over cellular networks. The mobile application communicated with "
        "the backend for authentication, sessions, and media persistence and with the ML API for inference, while IoT "
        "devices uploaded directly to authenticated backend endpoints. This architecture separated application logic, "
        "persistence, and inference responsibilities while maintaining HTTPS accessibility for field deployment.",
    )

    add_para(
        doc,
        f"The sixth specific objective—machine learning model evaluation—was achieved. Standard classification metrics "
        f"including accuracy, precision, recall, F1-score, confusion matrices, and training-curve compilations "
        f"demonstrated acceptable offline performance for a pre-screening prototype. Cough TB recall (47.63%) and "
        f"sputum specificity (50.00% on only six AFB-negative test samples) indicate that the models support triage "
        f"prioritization rather than standalone clinical diagnosis. Classification response time was largely within "
        f"the five-second UAT target under stable connectivity, with occasional delays reported on unstable networks.",
    )

    add_para(
        doc,
        f"The seventh and eighth specific objectives—ISO/IEC 25010 software quality assessment and professional user "
        f"acceptance testing—were achieved. Six evaluators representing IT, nursing, medical laboratory science, and "
        f"clinical instruction roles completed six predefined screening tasks and a twenty-four-item questionnaire "
        f"mapped to Functional Suitability, Performance Efficiency, Usability, Reliability, Security, and Overall "
        f"Satisfaction. The ISO/IEC 25010 grand mean was {grand_mean:.2f} ({grand_interp}), with Functional "
        f"Suitability ({uat['dimensions']['functional']['mean']:.2f}), Usability ({uat['dimensions']['usability']['mean']:.2f}), "
        f"Security ({uat['dimensions']['security']['mean']:.2f}), and Overall Satisfaction "
        f"({uat['dimensions']['satisfaction']['mean']:.2f}) rated Strongly Agree. Performance Efficiency "
        f"({uat['dimensions']['performance']['mean']:.2f}) and Reliability ({uat['dimensions']['reliability']['mean']:.2f}) "
        f"were rated Agree, reflecting connectivity and rural-deployment concerns raised during evaluation. Overall "
        f"UAT performance ratings were Excellent (33.3%), Good (50.0%), and Fair (16.7%)."
        if uat
        else
        f"The seventh and eighth specific objectives—ISO/IEC 25010 software quality assessment and professional user "
        f"acceptance testing—were achieved through structured evaluation with six professional evaluators.",
    )

    add_para(
        doc,
        "Functional and inference test cases executed on iPhone 15 and Redmi Note 12 Pro 5G devices confirmed that "
        "end-to-end screening workflow cases (TC-01 through TC-06) passed on both platforms and that cough quality "
        "gating correctly rejected noise, silence, and replay-like inputs while returning consistent API outputs for "
        "valid cough bursts and held-out sputum smear images. These results support the practical deployability of the "
        "integrated prototype within booth-operated TB pre-screening workflows.",
    )

    add_para(
        doc,
        "The utilization of CRISP-DM methodology, agile sprint development, multimodal machine learning, cloud "
        "infrastructure, and IoT-assisted capture demonstrated the applicability of engineering-oriented health "
        "informatics systems within resource-limited Philippine screening contexts. TBhon was consistently positioned "
        "and evaluated as pre-screening and triage support only—not a substitute for GeneXpert, smear microscopy, "
        "chest radiography, or clinician-confirmed TB diagnosis. Evaluation was conducted on publicly available "
        "datasets and prototype booth conditions; therefore, results represent technical feasibility and user "
        "acceptability rather than formal clinical validation.",
    )

    add_para(
        doc,
        "Overall, the proposed TBhon system successfully achieved the research objectives by providing a practical, "
        "accessible, and emerging-technology-driven multimodal TB pre-screening platform that combines symptom "
        "checklist assessment, cough and sputum machine learning signals, IoT-assisted capture, cloud-backed session "
        "management, and staff-reviewed triage outputs suitable for community health worker and booth deployment "
        "contexts in support of earlier referral and organized TB screening activities.",
    )

    # --- 7.2 Recommendations ---
    add_heading(doc, "7.2 Recommendations", 2)
    add_para(
        doc,
        "Based on the findings, limitations, and user acceptance feedback documented in Chapter VI, the researchers "
        "recommend the following improvements for future development, field pilot deployment, and research extension.",
    )

    recommendations = [
        (
            "First,",
            "future developers should implement offline or low-bandwidth operation with local request queuing and "
            "synchronization when connectivity is restored. UAT evaluators consistently identified unstable Wi-Fi and "
            "rural network conditions as the primary limitation (Item 20, weighted mean 3.83), and this enhancement "
            "would improve reliability for barangay health stations and remote screening booths.",
        ),
        (
            "Second,",
            "the mobile application should include a structured onboarding tutorial, illustrated capture guidance, "
            "and clearer ML processing status indicators for first-time booth staff. Evaluators from clinical and "
            "medical laboratory science backgrounds noted that non-technical users require additional instructional "
            "support before independent operation.",
        ),
        (
            "Third,",
            "referral outputs should integrate location-based mapping to the nearest accredited TB DOTS clinic, RHU, "
            "or GeneXpert-capable facility. Clinical evaluators recommended that triage results explicitly link "
            "patients to geographically relevant follow-up pathways rather than generic referral text alone.",
        ),
        (
            "Fourth,",
            "the cough ML pipeline should be expanded with additional training data—including female cough recordings, "
            "diverse acoustic environments, and improved quality-gate feedback that highlights rejected segments rather "
            "than full replay loops—to improve TB-class recall and user trust in audio-based triage signals.",
        ),
        (
            "Fifth,",
            "the sputum analysis module should be enhanced with on-screen AFB bounding-box visualization, expanded "
            "AFB-negative training samples, and prospective field validation using institution-approved clinical smear "
            "images to strengthen specificity estimates beyond the limited six-sample negative test partition.",
        ),
        (
            "Sixth,",
            "the symptom checklist may be extended with additional clinical risk factors suggested during UAT—such as "
            "comorbid conditions, smoking exposure, and allergy-related cough differentiators—while preserving a "
            "concise staff-operated workflow suitable for booth throughput.",
        ),
        (
            "Seventh,",
            "IoT deployment should be completed through native or EAS builds enabling full BLE Wi-Fi provisioning, "
            "on-device status LEDs or display feedback, and field troubleshooting guides to reduce setup friction for "
            "ESP32-S3 cough and sputum capture modules.",
        ),
        (
            "Eighth,",
            "a multi-site field pilot with LGU, RHU, or NGO partners should be conducted to measure real screening "
            "volume, referral follow-through, cloud inference latency under production load, and long-term model "
            "drift using logged prediction data from the deployed FastAPI service.",
        ),
        (
            "Ninth,",
            "future researchers may integrate supplementary modalities—such as chest X-ray CAD interfaces or "
            "GeneXpert result logging—as optional triage inputs only, maintaining honest pre-screening scope and "
            "avoiding presentation of the fused risk score as a standalone diagnostic conclusion.",
        ),
        (
            "Finally,",
            "multilingual support (e.g., Filipino and Cebuano/Bisaya) for checklist items, Learn content, result "
            "counseling, and disclaimers is recommended to improve accessibility for diverse clients in Mindanao "
            "screening environments and to align with community-oriented tuberculosis elimination initiatives.",
        ),
    ]
    for lead, body in recommendations:
        add_recommendation(doc, lead, body)

    add_para(
        doc,
        "The researchers strongly recommend the continuous exploration and integration of emerging technologies "
        "within community TB pre-screening and public health monitoring environments to improve multimodal triage "
        "efficiency, booth-operated screening capability, and technology-assisted tuberculosis elimination "
        "sustainability in resource-limited settings.",
    )

    try:
        doc.save(OUT)
        print(f"Wrote {OUT}")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"Wrote {OUT_FALLBACK} (close {OUT.name} in Word to overwrite primary file)")


if __name__ == "__main__":
    build()
