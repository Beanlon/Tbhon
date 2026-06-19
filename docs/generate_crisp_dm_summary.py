"""Generate TBhon CRISP-DM Progress and Output Summary (Word)."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "TBhon_CRISP_DM_Progress_Summary.docx"

COUGH_METRICS = ROOT / "ml" / "runs" / "20260531_014419" / "metrics.json"
SPUTUM_METRICS = ROOT / "ml (phlegm)" / "runs" / "phlegm_afb_binary_20260531_133949" / "metrics.json"
PRODUCTION_MANIFEST = ROOT / "ml" / "production_model.json"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def pct(value: float, digits: int = 2) -> str:
    return f"{100.0 * value:.{digits}f}%"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    doc.add_paragraph()


def build() -> None:
    cough = load_json(COUGH_METRICS)
    sputum = load_json(SPUTUM_METRICS)
    production = load_json(PRODUCTION_MANIFEST) if PRODUCTION_MANIFEST.is_file() else {}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("TBhon — CRISP-DM Progress and Output Summary")
    t_run.bold = True
    t_run.font.size = Pt(14)
    t_run.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run(
        "An IoT-Enabled Mobile-Based Multimodal Diagnostic Support System "
        "for Tuberculosis Pre-screening"
    )
    s_run.font.size = Pt(12)
    s_run.font.name = "Times New Roman"
    doc.add_paragraph()

    add_para(
        doc,
        "This document summarizes the Cross-Industry Standard Process for Data Mining (CRISP-DM) "
        "progress and deliverables for the TBhon project. The overall system design follows the "
        "User-Centered System Design (UCSD) methodology documented in Chapter II; CRISP-DM applies "
        "specifically to the data mining and machine learning pipeline (datasets, preprocessing, "
        "model development, evaluation, and ML deployment).",
    )

    add_heading(doc, "Overview", 1)
    add_table(
        doc,
        ["CRISP-DM Phase", "Status", "Summary"],
        [
            [
                "1. Business Understanding",
                "Complete",
                "TB triage problem defined; screening-only scope; eight research objectives set",
            ],
            [
                "2. Data Understanding",
                "Complete",
                "Two public datasets explored; class imbalance and labeling workflow documented",
            ],
            [
                "3. Data Preparation",
                "Complete",
                "Audio and image preprocessing pipelines implemented and reproducible",
            ],
            [
                "4. Modeling",
                "Complete",
                "Hybrid cough model, ResNet18 sputum model, and rule-based multimodal fusion",
            ],
            [
                "5. Evaluation",
                "Substantially complete",
                "Held-out ML metrics, functional test cases, UAT (4/5 evaluators), ISO 25010 pending final evaluator",
            ],
            [
                "6. Deployment",
                "Substantially complete",
                "Cloud ML API live; mobile E2E verified; IoT bench-tested; BLE provisioning needs native build",
            ],
        ],
    )

    # --- Phase 1 ---
    add_heading(doc, "Phase 1 — Business Understanding", 1)
    add_para(doc, "Goal", bold=True)
    add_para(
        doc,
        "Address fragmented, single-modality tuberculosis screening in resource-limited settings "
        "(rural health units, screening booths, and barangay health stations) by building a "
        "staff-operated multimodal triage support system. TBhon is explicitly scoped as screening "
        "support—not a clinical diagnosis tool.",
    )
    add_para(doc, "Key Activities Completed", bold=True)
    add_bullets(
        doc,
        [
            "Problem statement and stakeholder analysis (community health workers, booth staff, patients, LGUs)",
            "Hierarchical Task Analysis: intake → checklist → cough (×3) → sputum → ML → staff review → result/QR",
            "Requirements gathering through interviews, survey/questionnaire, and observation",
            "Scope boundaries: in-scope triage risk (Low/Moderate/High), referral guidance, session history; "
            "out-of-scope treatment plans, clinician dashboards, and formal diagnosis",
        ],
    )
    add_para(doc, "Outputs", bold=True)
    add_table(
        doc,
        ["Output", "Evidence"],
        [
            ["Problem statement and assumptions", "docs/TBhon_Chapters_1-3.docx"],
            ["Functional and data requirements", "Chapter II, Requirements Gathering"],
            ["Eight specific research objectives", "Chapter VI, Section 6.1"],
            ["UAT criteria (ISO/IEC 25010 aligned)", "docs/TBhon_UAT_Form.docx"],
        ],
    )
    add_para(doc, "Business Success Criteria", bold=True)
    add_bullets(
        doc,
        [
            "Booth staff can complete the full screening workflow on mobile",
            "ML produces interpretable cough and sputum signals",
            "Fused risk supports triage prioritization",
            "System clearly disclaims screening-support-only scope",
        ],
    )

    # --- Phase 2 ---
    add_heading(doc, "Phase 2 — Data Understanding", 1)
    add_para(doc, "Datasets Identified and Characterized", bold=True)
    add_table(
        doc,
        ["Modality", "Dataset", "Label Type", "Notes"],
        [
            [
                "Cough audio",
                "Kaggle ruchikashirsath/tb-audio",
                "Binary: TB / no-TB",
                "Predefined fold CSVs; production test fold n = 2,606",
            ],
            [
                "Sputum microscopy",
                "Raw_Sputum_Microscopy_Dataset",
                "YOLO boxes → AFB count → binary or load grade",
                "Imbalanced; production test split n = 216 (~97% AFB-positive)",
            ],
            [
                "Checklist",
                "11 symptom/exposure questions",
                "Rule-based (not dataset-trained)",
                "Hand-crafted log-odds scoring for fusion",
            ],
        ],
    )
    add_para(doc, "Data Understanding Findings", bold=True)
    add_bullets(
        doc,
        [
            "Cough: class skew (~70% no-TB in test fold); TB class is harder to detect",
            "Sputum: few AFB-negative samples (6 in production test set); high accuracy can be misleading",
            "Quality issues: silence, speech, and steady noise in cough; invalid smear photos in sputum",
        ],
    )
    add_para(doc, "Outputs", bold=True)
    add_table(
        doc,
        ["Output", "Evidence"],
        [
            ["Sputum annotation workflow", "docs/figures/figure_5_3_sputum_annotation_workflow.png"],
            ["Dataset split strategy", "ml/train_tb_cough_cnn.py, ml (phlegm)/train_phlegm_cnn.py"],
            ["Class distribution in metrics", "metrics.json per training run"],
        ],
    )

    # --- Phase 3 ---
    add_heading(doc, "Phase 3 — Data Preparation", 1)
    add_para(doc, "Cough Audio Pipeline", bold=True)
    add_para(
        doc,
        "WAV → resample 16 kHz → clip/pad 4 s → log-Mel spectrogram (CNN branch); "
        "parallel extraction of MFCC, deltas, mel statistics, RMS, and zero-crossing rate (GBM branch).",
    )
    add_bullets(
        doc,
        [
            "Train/validation/test split from fold CSVs; 12% validation from train (stratified)",
            "Augmentation: noise, time shift, spectrogram masking, codec/reverb simulation",
        ],
    )
    add_para(doc, "Sputum Image Pipeline", bold=True)
    add_para(
        doc,
        "Microscopy image → count YOLO bounding boxes → binary AFB-positive/negative (or load grade) → "
        "resize 224×224 → ImageNet normalization → augmentation (flip, color jitter).",
    )
    add_bullets(
        doc,
        [
            "Stratified 70/15/15 resplit for binary AFB task",
            "Class-weighted sampling to mitigate imbalance",
        ],
    )
    add_para(doc, "Quality Gates (Pre-Model Filtering)", bold=True)
    add_table(
        doc,
        ["Gate", "Purpose", "Module"],
        [
            ["Cough authenticity check", "Reject silence, speech, replay, steady noise", "ml/cough_quality.py"],
            ["Sputum image QC", "Reject non-smear or unusable images", "ml/infer_api.py"],
        ],
    )
    add_para(doc, "Outputs", bold=True)
    add_bullets(
        doc,
        [
            "Preprocessing code: train_tb_cough_cnn.py, train_tb_cough_hybrid.py, train_phlegm_cnn.py",
            "Curated partitions saved per run under ml/runs/ and ml (phlegm)/runs/",
        ],
    )

    # --- Phase 4 ---
    add_heading(doc, "Phase 4 — Modeling", 1)
    add_para(doc, "Model 1 — Hybrid Cough Classifier (Production)", bold=True)
    add_table(
        doc,
        ["Component", "Method", "Role"],
        [
            ["CNN branch", "LegacySmallAudioCNN on log-Mel spectrograms", "Deep pattern learning"],
            ["GBM branch", "Gradient Boosting on hand-crafted features", "Classical audio statistics"],
            ["Fusion", "Weighted probability blend + calibrated threshold", "Final P(TB)"],
        ],
    )
    add_para(
        doc,
        f"Production run: {production.get('run_id', '20260531_014419')} "
        f"(ml/production_model.json). CNN trained 8 epochs; blend weight w = "
        f"{cough.get('blend_cnn_weight', 0.05)}; decision threshold = "
        f"{cough.get('decision_threshold', 0.395)}; selection metric = validation macro-F1.",
    )
    add_para(doc, "Model 2 — Sputum AFB Binary Classifier (Production)", bold=True)
    add_table(
        doc,
        ["Component", "Method"],
        [
            ["Backbone", "ResNet18 (ImageNet pretrained)"],
            ["Task", "afb_negative vs afb_positive"],
            ["Threshold policy", "Max specificity with sensitivity ≥ 95% on validation"],
        ],
    )
    add_para(
        doc,
        "Production run: phlegm_afb_binary_20260531_133949. Decision threshold = "
        f"{sputum.get('decision_threshold', 0.63)}.",
    )
    add_para(doc, "Model 3 — Multimodal Risk Fusion (Rule-Based)", bold=True)
    add_table(
        doc,
        ["Input", "Method"],
        [
            ["Checklist", "Logistic symptom model (hand-tuned log-odds)"],
            ["Cough", "Hybrid model prob_tb"],
            ["Sputum", "CNN prob(afb_positive)"],
            ["Combine", "Weighted log-odds (0.85 / 1.0 / 0.7) + clinical safety floors"],
        ],
    )
    add_para(doc, "Implementation: ml/tb_risk_fusion.py, mobile/utils/tbRiskFusion.ts", bold=False)
    add_para(doc, "Outputs", bold=True)
    add_bullets(
        doc,
        [
            "Trained checkpoints: model.pt, hybrid_bundle.pkl, phlegm model_last.pt",
            "26+ experiment runs under ml/runs/ and ml (phlegm)/runs/",
            "Inference: hybrid_predict.py, infer_phlegm.py, infer_api.py",
            "Architecture figures in docs/figures/",
        ],
    )

    # --- Phase 5 ---
    add_heading(doc, "Phase 5 — Evaluation", 1)
    add_para(doc, "5A — ML Model Evaluation (Held-Out Test Sets)", bold=True)
    add_para(doc, f"Cough hybrid classifier (n = 2,606)", bold=True)
    add_table(
        doc,
        ["Metric", "Result"],
        [
            ["Test accuracy", pct(cough["test_accuracy"])],
            ["Macro F1", pct(cough["best_f1_macro"])],
            ["TB recall (class 1)", "47.63%"],
            ["Non-TB recall (class 0)", "88.25%"],
        ],
    )
    add_para(doc, f"Sputum binary ResNet18 (n = 216)", bold=True)
    add_table(
        doc,
        ["Metric", "Result"],
        [
            ["Test accuracy", pct(sputum["test_acc"])],
            ["Macro F1", pct(sputum["test_macro_f1"])],
            ["AFB-positive sensitivity", pct(sputum["sensitivity"])],
            ["AFB-negative specificity", pct(sputum["specificity"])],
        ],
    )
    add_para(doc, "5B — Functional and Integration Test Cases", bold=True)
    add_table(
        doc,
        ["Test Type", "Cases", "Result"],
        [
            [
                "Cough QC",
                "Valid cough, noise, silence, speech",
                "QC accepts/rejects correctly; bad audio excluded from fusion",
            ],
            [
                "Sputum inference",
                "Four held-out test images",
                "AFB+/− predictions with confidence scores",
            ],
            [
                "End-to-end app",
                "Full screening workflow",
                "Verified per sprint report (cough + sputum + risk fusion)",
            ],
        ],
    )
    add_para(doc, "5C — Software Quality and UAT", bold=True)
    add_table(
        doc,
        ["Evaluation", "Status", "Notes"],
        [
            ["ISO/IEC 25010 instrument", "Done", "24 criteria across five characteristics"],
            ["UAT with professionals", "4 of 5 complete", "CHWs, nurses, IT; one med tech pending"],
            ["Heuristic evaluation (prototype)", "Done", "Nielsen heuristics, Chapter II"],
            ["Response time target", "≤ 5 s per modality", "Observed during E2E testing"],
        ],
    )
    add_para(doc, "Evaluation Limitations", bold=True)
    add_bullets(
        doc,
        [
            "Metrics are offline benchmarks on public datasets, not clinical validation",
            "Sputum test set contains only six AFB-negative samples",
            "Threshold tuning on validation introduces minor optimism bias",
            "Fused risk score was not separately validated as one combined model",
        ],
    )
    add_para(doc, "Outputs", bold=True)
    add_bullets(
        doc,
        [
            "docs/TBhon_ML_Validation_Metrics.docx",
            "docs/TBhon_Chapter_6.docx (Sections 6.2–6.7)",
            "docs/test_case_results.json",
            "Training-curve figures in docs/figures/",
        ],
    )

    # --- Phase 6 ---
    add_heading(doc, "Phase 6 — Deployment", 1)
    add_para(doc, "Deployed Architecture", bold=True)
    add_para(
        doc,
        "Expo mobile app → Node.js/Express/Prisma/MySQL backend (sessions, media, auth) → "
        "Python FastAPI/PyTorch ML API (/predict, /predict-phlegm, /fuse-risk). "
        "ESP32 IoT devices upload to backend /iot/* endpoints. "
        "Hosting: DigitalOcean droplets with Cloudflare HTTPS tunnels.",
    )
    add_para(doc, "Deployment Status", bold=True)
    add_table(
        doc,
        ["Component", "Status", "Evidence"],
        [
            ["ML FastAPI service", "Deployed", "ml/infer_api.py; systemd on ML droplet"],
            [
                "Production cough model",
                "Pinned",
                f"ml/production_model.json → run {production.get('run_id', '20260531_014419')}",
            ],
            ["Mobile ↔ ML integration", "E2E verified", "mobile/app/screening/processing.tsx"],
            ["Backend + database", "Deployed", "Prisma schema; DigitalOcean Managed MySQL"],
            ["IoT firmware upload", "Bench verified", "ESP32 hardcoded Wi-Fi for testing"],
            ["BLE Wi-Fi provisioning", "In progress", "Requires native/EAS build (not Expo Go)"],
            ["Model monitoring", "Basic", "Prediction logging and drift alerts in infer_api.py"],
        ],
    )
    add_para(doc, "Outputs", bold=True)
    add_bullets(
        doc,
        [
            "Deployment workflow figures (PyTorch/FastAPI, cloud backend)",
            "API documentation: Swagger /docs on backend and ML service",
            "ml/smoke_test_infer.py",
            "Tbhon-Backend/docs/sprint-report.md",
        ],
    )

    # --- Mapping ---
    add_heading(doc, "CRISP-DM ↔ Research Objectives Mapping", 1)
    add_table(
        doc,
        ["Research Objective", "CRISP-DM Phase(s)", "Status"],
        [
            ["1. IoT acquisition hardware", "Business Understanding + Deployment", "Substantially achieved"],
            ["2. Data curation and preprocessing", "Data Understanding + Data Preparation", "Achieved"],
            ["3. ML model development", "Modeling", "Achieved"],
            ["4. Multimodal risk scoring", "Modeling + Evaluation", "Achieved"],
            ["5. Communication architecture", "Deployment", "Achieved"],
            ["6. ML performance evaluation", "Evaluation", "Achieved"],
            ["7. ISO/IEC 25010 quality assessment", "Evaluation", "Substantially achieved"],
            ["8. Professional UAT", "Evaluation", "Substantially achieved (4/5)"],
        ],
    )

    add_heading(doc, "Executive Summary", 1)
    add_para(
        doc,
        f"TBhon followed the CRISP-DM process to develop a multimodal tuberculosis screening support "
        f"system. Business understanding established the need for staff-operated, non-diagnostic triage "
        f"in resource-limited Philippine health settings. Public cough audio and sputum microscopy datasets "
        f"were explored, preprocessed, and partitioned into held-out test sets. Modeling produced a hybrid "
        f"CNN+GBM cough classifier ({pct(cough['test_accuracy'])} test accuracy, "
        f"{pct(cough['best_f1_macro'])} macro F1) and a ResNet18 binary sputum classifier "
        f"({pct(sputum['test_acc'])} test accuracy, {pct(sputum['sensitivity'])} AFB-positive sensitivity), "
        f"combined with a rule-based weighted log-odds fusion layer for Low, Moderate, and High risk "
        f"stratification. Evaluation included standard classification metrics, functional test cases, "
        f"heuristic prototype review, and ISO/IEC 25010–aligned user acceptance testing (four of five "
        f"professional evaluators completed). Deployment integrated both models into a FastAPI inference "
        f"service on a cloud ML droplet, connected to an Expo mobile application and Node.js backend with "
        f"verified end-to-end screening workflow. Remaining gaps include final UAT evaluator completion, "
        f"native-build IoT BLE provisioning, and formal clinical validation beyond public-dataset benchmarks.",
    )

    add_para(doc, "")
    add_para(
        doc,
        "Document generated from project artifacts and production metrics. "
        "Regenerate with: python docs/generate_crisp_dm_summary.py",
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
