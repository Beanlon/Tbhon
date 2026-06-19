"""Generate TBhon Appendix F — Model Performance Outputs (Word)."""
from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "docs" / "figures"
OUT = ROOT / "docs" / "TBhon_Appendix_F.docx"
OUT_FALLBACK = ROOT / "docs" / "TBhon_Appendix_F_generated.docx"

COUGH_METRICS = ROOT / "ml" / "runs" / "20260531_014419" / "metrics.json"
COUGH_LOG = ROOT / "ml" / "runs" / "20260531_014419" / "epoch_log.jsonl"
SPUTUM_METRICS = ROOT / "ml (phlegm)" / "runs" / "phlegm_afb_binary_20260531_133949" / "metrics.json"
FIG_SCRIPT = ROOT / "docs" / "generate_appendix_f_figures.py"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_figure(doc: Document, path: Path, width: float = 5.8) -> None:
    if path.is_file():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"[Insert figure: {path.name}]")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    doc.add_paragraph()


def pct(v: float) -> str:
    return f"{100.0 * v:.2f}%"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_figures() -> None:
    if not (FIG / "appendix_f_1_precision_test.png").is_file():
        subprocess.run([sys.executable, str(FIG_SCRIPT)], check=True, cwd=str(ROOT))


def build() -> None:
    ensure_figures()
    cough = load_json(COUGH_METRICS)
    sputum = load_json(SPUTUM_METRICS)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    add_heading(doc, "APPENDIX F", 1)
    add_heading(doc, "MODEL PERFORMANCE OUTPUTS", 1)
    add_para(
        doc,
        "This appendix presents supplementary model performance evidence supporting the Chapter VI evaluation "
        "results for the TBhon hybrid cough classifier (production run 20260531_014419, fold 1) and the sputum "
        "AFB binary ResNet18 classifier (production run phlegm_afb_binary_20260531_133949). Figures include "
        "test-set precision and recall summaries, validation macro F1 curves, training loss curves, confusion "
        "matrices, and consolidated training summary outputs.",
    )

    add_heading(doc, "F.1 Precision Results", 2)
    add_para(
        doc,
        "Figure F.1 summarizes held-out test-set precision for the non-TB/TB cough classes and the AFB-negative/"
        "AFB-positive sputum classes. Cough non-TB precision was 79.13% and TB precision was 64.31%. Sputum "
        "AFB-positive precision was 98.54%; AFB-negative precision was 27.27% (n = 6 negatives in the test partition).",
    )
    add_figure(doc, FIG / "appendix_f_1_precision_test.png")
    add_caption(doc, "Figure F.1. Test-Set Precision by Class (Cough and Sputum Models)")

    add_heading(doc, "F.2 Recall Results", 2)
    add_para(
        doc,
        "Figure F.2 summarizes held-out test-set recall. Cough non-TB recall was 88.25% and TB recall was 47.63%. "
        "Sputum AFB-positive recall (sensitivity) was 96.19% under the validation-tuned threshold policy "
        f"(decision threshold = {sputum['decision_threshold']}).",
    )
    add_figure(doc, FIG / "appendix_f_2_recall_test.png")
    add_caption(doc, "Figure F.2. Test-Set Recall by Class (Cough and Sputum Models)")

    add_heading(doc, "F.3 Validation Macro F1 Curves", 2)
    add_para(
        doc,
        "Figure F.3 presents validation macro F1 curves across training epochs for the cough CNN branch (8 epochs) "
        "and the sputum ResNet18 classifier (30 epochs). Macro F1 is used as the primary validation metric for "
        "binary classification in place of mean average precision (mAP), which applies to object-detection tasks.",
    )
    add_figure(doc, FIG / "appendix_f_3_macro_f1_curves.png", 6.2)
    add_caption(doc, "Figure F.3. Validation Macro F1 Curves")

    add_heading(doc, "F.4 Training Loss Curves", 2)
    add_para(
        doc,
        "Figure F.4 shows training loss curves for both production models. Decreasing cough CNN training loss and "
        "sputum ResNet18 training loss indicate improved fit during the reported training runs.",
    )
    add_figure(doc, FIG / "appendix_f_4_train_loss_curves.png", 6.2)
    add_caption(doc, "Figure F.4. Training Loss Curves")

    add_heading(doc, "F.5 Confusion Matrices", 2)
    add_para(
        doc,
        "Figure F.5 presents test-set confusion matrices for the hybrid cough classifier (n = 2,606) and the binary "
        "sputum AFB classifier (n = 216). Cough matrix: TN 1592, FP 212, FN 420, TP 382. Sputum matrix: TN 3, FP 3, "
        "FN 8, TP 202.",
    )
    add_figure(doc, FIG / "appendix_f_5_confusion_matrices.png", 6.2)
    add_caption(doc, "Figure F.5. Test-Set Confusion Matrices")

    add_heading(doc, "F.6 Training Summary Output", 2)
    add_para(doc, "Table F.6.1. Cough Hybrid Classifier — Production Test Summary (run 20260531_014419)")
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Test accuracy", pct(cough["test_accuracy"])],
            ["Macro F1", pct(cough["best_f1_macro"])],
            ["Non-TB precision", "79.13%"],
            ["TB precision", "64.31%"],
            ["Non-TB recall", "88.25%"],
            ["TB recall", "47.63%"],
            ["CNN blend weight", str(cough.get("blend_cnn_weight", "—"))],
            ["Decision threshold", str(cough.get("decision_threshold", "—"))],
            ["Test samples (n)", "2,606"],
        ],
    )

    add_para(doc, "Table F.6.2. Sputum ResNet18 AFB Binary Classifier — Production Test Summary (phlegm_afb_binary_20260531_133949)")
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Test accuracy", pct(sputum["test_acc"])],
            ["Macro F1", pct(sputum["test_macro_f1"])],
            ["AFB+ sensitivity (recall)", pct(sputum["sensitivity"])],
            ["AFB− specificity", pct(sputum["specificity"])],
            ["AFB+ precision", "98.54%"],
            ["Decision threshold", str(sputum["decision_threshold"])],
            ["Min. sensitivity constraint", pct(sputum.get("min_sensitivity_constraint", 0.95))],
            ["Test samples (n)", "216"],
        ],
    )

    add_para(doc, "Table F.6.3. Cough CNN Branch — Per-Epoch Validation Log (8 epochs)")
    if COUGH_LOG.is_file():
        epoch_rows = []
        for line in COUGH_LOG.read_text(encoding="utf-8").splitlines():
            if not line.strip():
                continue
            row = json.loads(line)
            epoch_rows.append(
                [
                    str(row["epoch"]),
                    f"{row['train_loss']:.4f}",
                    f"{row['val_accuracy']:.2%}",
                    f"{row['val_f1_macro']:.2%}",
                ]
            )
        add_table(doc, ["Epoch", "Train Loss", "Val Accuracy", "Val Macro F1"], epoch_rows)

    add_para(doc, "Table F.6.4. Cough Hybrid Classifier — Per-Class Test Metrics")
    add_table(
        doc,
        ["Class", "Precision", "Recall", "F1-Score", "Support"],
        [
            ["Non-TB (0)", "79.13%", "88.25%", "83.44%", "1,804"],
            ["TB (1)", "64.31%", "47.63%", "54.73%", "802"],
            ["Macro average", "71.72%", "67.94%", "69.08%", "2,606"],
            ["Weighted average", "74.57%", "75.75%", "74.60%", "2,606"],
        ],
    )

    add_para(doc, "Table F.6.5. Sputum AFB Binary Classifier — Per-Class Test Metrics")
    add_table(
        doc,
        ["Class", "Precision", "Recall", "F1-Score", "Support"],
        [
            ["AFB-negative", "27.27%", "50.00%", "35.29%", "6"],
            ["AFB-positive", "98.54%", "96.19%", "97.11%", "210"],
            ["Macro average", "62.91%", "73.10%", "66.32%", "216"],
            ["Weighted average", "96.85%", "94.91%", "95.85%", "216"],
        ],
    )

    try:
        doc.save(OUT)
        print(f"Wrote {OUT}")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"Wrote {OUT_FALLBACK} (close {OUT.name} in Word to overwrite primary file)")


if __name__ == "__main__":
    build()
